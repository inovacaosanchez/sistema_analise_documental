"""
LEITOR DE ARQUIVOS
Responsável por ler e extrair conteúdo de documentos Word
"""

import os
from pathlib import Path
from docx import Document
import re
from typing import Dict, List, Tuple, Optional

from utils.logger import SystemLogger
from config.settings import Settings

class FileReader:
    """Leitor especializado para documentos Word"""
    
    def __init__(self, logger: SystemLogger):
        self.logger = logger
        self.settings = Settings()
    
    def scan_directory_structure(self, root_path: str) -> List[Dict]:
        """
        Escaneia estrutura de diretórios procurando documentos
        Retorna lista de documentos encontrados com metadados
        """
        root_path = Path(root_path)
        documents = []
        
        if not root_path.exists():
            self.logger.error(f"Caminho não encontrado: {root_path}")
            return documents
        
        self.logger.info(f"🔍 Escaneando estrutura: {root_path}")
        
        try:
            # Percorrer estrutura DEPARTAMENTO/SUBAREA/
            for dept_path in root_path.iterdir():
                if not dept_path.is_dir():
                    continue
                
                departamento = dept_path.name
                self.logger.debug(f"📁 Departamento: {departamento}")
                
                for subarea_path in dept_path.iterdir():
                    if not subarea_path.is_dir():
                        continue
                    
                    subarea = subarea_path.name
                    self.logger.debug(f"📂 Subárea: {subarea}")
                    
                    # Procurar documentos Word na subárea
                    for file_path in subarea_path.iterdir():
                        if (file_path.is_file() and 
                            file_path.suffix.lower() in self.settings.SUPPORTED_EXTENSIONS):
                            
                            doc_info = {
                                'file_path': str(file_path),
                                'departamento': departamento,
                                'subarea': subarea,
                                'filename': file_path.stem,
                                'extension': file_path.suffix
                            }
                            
                            documents.append(doc_info)
                            self.logger.info(f"📄 Documento encontrado: {file_path.name}")
            
            self.logger.info(f"✅ Escaneamento concluído: {len(documents)} documentos encontrados")
            return documents
            
        except Exception as e:
            self.logger.error(f"Erro ao escanear diretório: {e}")
            return documents
    
    def read_document_content(self, file_path: str) -> Tuple[str, str, str]:
        """
        Lê conteúdo de documento Word
        Retorna: (titulo_extraido, texto_completo, texto_normalizado)
        """
        try:
            self.logger.info(f"📖 Lendo documento: {Path(file_path).name}")
            
            doc = Document(file_path)
            
            # Extrair título (primeiro parágrafo com estilo de título ou primeiro parágrafo não vazio)
            titulo = self._extract_title(doc)
            
            # Extrair texto completo
            texto_completo = self._extract_full_text(doc)
            
            # Normalizar texto
            texto_normalizado = self._normalize_text(texto_completo)
            
            # Validar tamanho mínimo
            if len(texto_normalizado) < self.settings.MIN_DOCUMENT_LENGTH:
                self.logger.warning(f"⚠️ Documento muito pequeno: {len(texto_normalizado)} caracteres")
            
            self.logger.info(f"✅ Documento lido: {len(texto_normalizado)} caracteres")
            return titulo, texto_completo, texto_normalizado
            
        except Exception as e:
            self.logger.error(f"Erro ao ler documento {file_path}: {e}")
            return "", "", ""
    
    def _extract_title(self, doc: Document) -> str:
        """Extrai título do documento"""
        try:
            # Procurar por estilos de título
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Verificar se é estilo de título
                    if (paragraph.style.name.startswith('Heading') or 
                        paragraph.style.name.startswith('Title') or
                        'título' in paragraph.style.name.lower()):
                        return paragraph.text.strip()
            
            # Se não encontrou título, usar primeiro parágrafo não vazio
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text and len(text) > 5:  # Mínimo de 5 caracteres
                    return text
            
            return ""
            
        except Exception as e:
            self.logger.warning(f"Erro ao extrair título: {e}")
            return ""
    
    def _extract_full_text(self, doc: Document) -> str:
        """Extrai texto completo do documento"""
        try:
            paragraphs = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extrair texto de tabelas também
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            paragraphs.append(text)
            
            return '\n\n'.join(paragraphs)
            
        except Exception as e:
            self.logger.error(f"Erro ao extrair texto: {e}")
            return ""
    
    def _normalize_text(self, text: str) -> str:
        """Normaliza texto removendo ruídos"""
        try:
            # Remover múltiplas quebras de linha
            text = re.sub(r'\n{3,}', '\n\n', text)
            
            # Remover espaços excessivos
            text = re.sub(r' {2,}', ' ', text)
            
            # Remover tabs
            text = text.replace('\t', ' ')
            
            # Remover caracteres especiais problemáticos
            text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]"\'\/\n]', '', text)
            
            # Trim geral
            text = text.strip()
            
            return text
            
        except Exception as e:
            self.logger.warning(f"Erro ao normalizar texto: {e}")
            return text
    
    def check_output_files_exist(self, output_path: str) -> Tuple[bool, List[str]]:
        """
        Verifica se arquivos de saída já existem
        Retorna: (algum_existe, lista_arquivos_existentes)
        """
        output_path = Path(output_path)
        existing_files = []
        
        for filename in self.settings.OUTPUT_FILES:
            file_path = output_path / filename
            if file_path.exists():
                existing_files.append(filename)
        
        return len(existing_files) > 0, existing_files