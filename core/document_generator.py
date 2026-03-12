"""
GERADOR DE DOCUMENTOS WORD
Responsavel por criar arquivos .docx com estrutura padronizada.
"""

from datetime import datetime
from pathlib import Path
import re
from typing import Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from utils.logger import SystemLogger
from config.settings import Settings


class DocumentGenerator:
    """Gerador de documentos Word formatados e padronizados."""

    def __init__(self, logger: SystemLogger):
        self.logger = logger
        self.settings = Settings()

    def create_all_documents(self, content: Dict[str, str], output_path: str, processo_name: str) -> Dict[str, bool]:
        output_path = Path(output_path)
        results = {}

        try:
            output_path.mkdir(parents=True, exist_ok=True)
            self.logger.info(f"Diretorio criado: {output_path}")

            file_mapping = {
                "Ficha_Tecnica.docx": content.get("ficha_tecnica", ""),
                "Fluxograma_Mermaid.docx": content.get("fluxograma", ""),
                "Riscos_e_Automatizacoes.docx": content.get("riscos", ""),
            }

            for filename, doc_content in file_mapping.items():
                file_path = output_path / filename
                try:
                    self.logger.info(f"Criando: {filename}")
                    results[filename] = self._create_single_document(doc_content, file_path, filename, processo_name)
                except Exception as exc:
                    results[filename] = False
                    self.logger.error(f"Erro ao criar {filename}: {exc}")

            return results
        except Exception as exc:
            self.logger.error(f"Erro geral na criacao de documentos: {exc}")
            return {f: False for f in self.settings.OUTPUT_FILES}

    def _create_single_document(self, content: str, file_path: Path, filename: str, processo_name: str) -> bool:
        try:
            doc = Document()
            self._setup_document_styles(doc)
            self._add_document_header(doc, filename, processo_name)

            if "Fluxograma" in filename:
                self._process_mermaid_content(doc, content)
            else:
                self._process_markdown_content(doc, content)

            self._add_document_footer(doc)
            doc.save(str(file_path))
            self.logger.info(f"{filename} criado com sucesso")
            return True
        except Exception as exc:
            self.logger.error(f"Erro ao criar documento {filename}: {exc}")
            return False

    def _setup_document_styles(self, doc: Document):
        section = doc.sections[0]
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(11)

        doc.styles["Heading 1"].font.name = "Calibri"
        doc.styles["Heading 1"].font.size = Pt(16)
        doc.styles["Heading 1"].font.bold = True

        doc.styles["Heading 2"].font.name = "Calibri"
        doc.styles["Heading 2"].font.size = Pt(15)
        doc.styles["Heading 2"].font.bold = True

        doc.styles["Heading 3"].font.name = "Calibri"
        doc.styles["Heading 3"].font.size = Pt(14)
        doc.styles["Heading 3"].font.bold = True

    def _add_document_header(self, doc: Document, filename: str, processo_name: str):
        title_map = {
            "Ficha_Tecnica.docx": "FICHA TECNICA DO PROCESSO",
            "Fluxograma_Mermaid.docx": "FLUXOGRAMA DO PROCESSO",
            "Riscos_e_Automatizacoes.docx": "ANALISE DE RISCOS E AUTOMATIZACAO",
        }
        title = title_map.get(filename, "DOCUMENTO DE PROCESSO")

        p_title = doc.add_paragraph(title)
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_title.runs[0]
        run.font.name = "Calibri"
        run.font.size = Pt(16)
        run.font.bold = True

        p_proc = doc.add_paragraph(f"Processo: {processo_name}")
        p_proc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_proc.runs[0].font.size = Pt(11)

        sep = doc.add_paragraph("-" * 70)
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sep.runs[0].font.size = Pt(9)

        doc.add_paragraph()

    def _add_document_footer(self, doc: Document):
        doc.add_paragraph()
        sep = doc.add_paragraph("-" * 70)
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sep.runs[0].font.size = Pt(9)

        footer = doc.add_paragraph(f"Documento gerado automaticamente em {self._get_current_datetime()}")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.size = Pt(9)
        footer.runs[0].italic = True

    def _process_markdown_content(self, doc: Document, content: str):
        try:
            lines = content.splitlines()
            table_buffer: List[str] = []
            current_topic_paragraph = None

            for raw_line in lines:
                line = raw_line.strip()

                if self._is_table_line(line):
                    current_topic_paragraph = None
                    table_buffer.append(line)
                    continue

                if table_buffer:
                    self._add_markdown_table(doc, table_buffer)
                    table_buffer = []

                if not line:
                    current_topic_paragraph = None
                    continue

                if line.startswith("# "):
                    doc.add_paragraph(line[2:].strip(), style="Heading 1")
                    current_topic_paragraph = None
                    continue

                if line.startswith("## "):
                    doc.add_paragraph(line[3:].strip(), style="Heading 2")
                    current_topic_paragraph = None
                    continue

                if line.startswith("### "):
                    doc.add_paragraph(line[4:].strip(), style="Heading 3")
                    current_topic_paragraph = None
                    continue

                if re.match(r"^\d+[\.\)]\s+", line):
                    text = re.sub(r"^\d+[\.\)]\s+", "", line).strip()
                    self._add_rich_text_paragraph(doc, text, style="List Number")
                    current_topic_paragraph = None
                    continue

                if line.startswith("- ") or line.startswith("* "):
                    self._add_rich_text_paragraph(doc, line[2:].strip(), style="List Bullet")
                    current_topic_paragraph = None
                    continue

                # Mesmo topico: quebra de linha simples no mesmo paragrafo.
                if current_topic_paragraph is None:
                    current_topic_paragraph = doc.add_paragraph(style="Normal")
                    self._add_rich_text_to_paragraph(current_topic_paragraph, line)
                else:
                    br = current_topic_paragraph.add_run()
                    br.add_break()
                    self._add_rich_text_to_paragraph(current_topic_paragraph, line)

            if table_buffer:
                self._add_markdown_table(doc, table_buffer)
        except Exception as exc:
            self.logger.error(f"Erro ao processar markdown: {exc}")
            doc.add_paragraph(content)

    def _process_mermaid_content(self, doc: Document, content: str):
        try:
            doc.add_paragraph("Orientacoes de uso", style="Heading 2")
            doc.add_paragraph("Este documento contem o fluxograma em linguagem Mermaid.")
            doc.add_paragraph("Para visualizar, copie o codigo e cole em um editor Mermaid:")
            doc.add_paragraph("https://mermaid.live/", style="List Bullet")
            doc.add_paragraph("https://mermaid-js.github.io/mermaid-live-editor/", style="List Bullet")

            doc.add_paragraph("")
            doc.add_paragraph("Codigo Mermaid", style="Heading 2")
            mermaid_code = self._extract_mermaid_code(content)
            code_para = doc.add_paragraph(mermaid_code)
            for run in code_para.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(10)

            remaining_content = content.replace(mermaid_code, "").strip()
            if remaining_content:
                doc.add_paragraph("")
                doc.add_paragraph("Observacoes adicionais", style="Heading 2")
                self._process_markdown_content(doc, remaining_content)
        except Exception as exc:
            self.logger.error(f"Erro ao processar Mermaid: {exc}")
            doc.add_paragraph(content)

    def _extract_mermaid_code(self, content: str) -> str:
        try:
            mermaid_block = re.search(r"```mermaid\s*(.*?)```", content, re.DOTALL | re.IGNORECASE)
            if mermaid_block:
                return mermaid_block.group(1).strip()

            flowchart = re.search(r"(flowchart[\s\S]*?)(?:\n\s*\n|$)", content, re.IGNORECASE)
            if flowchart:
                return flowchart.group(1).strip()

            return content.strip()
        except Exception:
            return content

    def _is_table_line(self, line: str) -> bool:
        return "|" in line and line.count("|") >= 2

    def _add_markdown_table(self, doc: Document, table_lines: List[str]):
        rows = []
        for line in table_lines:
            line = line.strip()
            if not line:
                continue
            cols = [c.strip() for c in line.strip("|").split("|")]
            if cols:
                rows.append(cols)

        if not rows:
            return

        if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in rows[1]):
            rows.pop(1)

        col_count = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=col_count)
        table.style = "Table Grid"

        for r_idx, row in enumerate(rows):
            for c_idx in range(col_count):
                text = row[c_idx] if c_idx < len(row) else ""
                cell_para = table.cell(r_idx, c_idx).paragraphs[0]
                cell_para.text = text
                if r_idx == 0:
                    for run in cell_para.runs:
                        run.bold = True

        doc.add_paragraph("")

    def _add_rich_text_paragraph(self, doc: Document, text: str, style: str = "Normal"):
        para = doc.add_paragraph(style=style)
        self._add_rich_text_to_paragraph(para, text)

    def _add_rich_text_to_paragraph(self, para, text: str):
        chunks = re.split(r"(\*\*.*?\*\*)", text)
        if not chunks:
            para.add_run(text)
            return

        for chunk in chunks:
            if not chunk:
                continue
            if chunk.startswith("**") and chunk.endswith("**") and len(chunk) >= 4:
                run = para.add_run(chunk[2:-2])
                run.bold = True
            else:
                para.add_run(chunk)

    def _get_current_datetime(self) -> str:
        return datetime.now().strftime("%d/%m/%Y as %H:%M")
