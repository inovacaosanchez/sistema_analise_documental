"""
WEB APPLICATION LAYER
Interface HTTP for document analysis and generation workflow.
"""

from __future__ import annotations

import hashlib
import math
import os
import secrets
import time
import unicodedata
from pathlib import Path
from threading import Lock, RLock
import threading
from io import BytesIO
from uuid import uuid4
from datetime import datetime, timedelta
from typing import Dict, List, Optional, Any
import re

from flask import Flask, jsonify, render_template, request, send_file, session, redirect, url_for, g
from docx import Document

from config.settings import Settings
from core.chunking_engine import ChunkingEngine
from core.content_generator import ContentGenerator
from core.document_generator import DocumentGenerator
from core.file_reader import FileReader
from repositories import (
    JsonRepositoryStore,
    SupabaseRepositoryStore,
    ProjectRepository,
    ProcessRepository,
    UserRepository,
    RoleRepository,
    DepartmentRepository,
    CargoRepository,
    SectorRepository,
)
from utils.logger import SystemLogger


class WebApp:
    """Main web app coordinator."""

    def __init__(self):
        self.logger = SystemLogger()
        self.settings = Settings()

        self.file_reader = FileReader(self.logger)
        self.chunking_engine = ChunkingEngine(self.logger)
        self.content_generator = ContentGenerator(self.logger)
        self.document_generator = DocumentGenerator(self.logger)

        self.lock = Lock()
        self.origem_path = ""
        self.destino_path = ""
        self.documents: List[Dict] = []
        self.processing = False
        self.current_index = 0
        self.current_preview: Optional[Dict] = None
        self.pending_previews: Dict[int, Dict] = {}
        self.worker_thread: Optional[threading.Thread] = None
        self.api_status = {"ok": False, "message": "Nao testado"}
        self.data_dir = Path(__file__).resolve().parent.parent / "data"
        self.projects_file = self.data_dir / "projects.json"
        self.sectors_file = self.data_dir / "sectors.json"
        self.departments_file = self.data_dir / "departments.json"
        self.cargos_file = self.data_dir / "cargos.json"
        self.processes_file = self.data_dir / "processes_registry.json"
        self.users_file = self.data_dir / "users.json"
        self.roles_file = self.data_dir / "roles.json"
        self.repo_store = self._build_repository_store()
        self.project_repo = ProjectRepository(self.repo_store, self.projects_file)
        self.process_repo = ProcessRepository(self.repo_store, self.processes_file)
        self.user_repo = UserRepository(self.repo_store, self.users_file)
        self.role_repo = RoleRepository(self.repo_store, self.roles_file)
        self.department_repo = DepartmentRepository(self.repo_store, self.departments_file)
        self.cargo_repo = CargoRepository(self.repo_store, self.cargos_file)
        self.sector_repo = SectorRepository(self.repo_store, self.sectors_file)
        self.auth_salt = os.getenv("AUTH_SALT", "hpo-auth-salt")
        self.app_env = str(os.getenv("APP_ENV", "development")).strip().lower()
        self.is_production = self.app_env in {"prod", "production"}
        self.enable_security_headers = str(os.getenv("ENABLE_SECURITY_HEADERS", "1")).strip() != "0"
        self.enable_csrf = str(os.getenv("ENABLE_CSRF", "1")).strip() != "0"
        self.enable_login_rate_limit = str(os.getenv("ENABLE_LOGIN_RATE_LIMIT", "1")).strip() != "0"
        self.csrf_exempt_paths = {"/api/auth/login"}
        self.login_rate_limit_store: Dict[str, List[float]] = {}
        self.login_rate_lock = RLock()
        self.login_rate_window_sec = int(os.getenv("LOGIN_RATE_WINDOW_SEC", "300"))
        self.login_rate_ip_max = int(os.getenv("LOGIN_RATE_IP_MAX", "20"))
        self.login_rate_user_max = int(os.getenv("LOGIN_RATE_USER_MAX", "10"))
        self.available_views = [
            "view-processos-analise",
            "view-processos-priorizacao",
            "view-processos-cadastro",
            "view-processos-consulta",
            "view-processos-dashboard",
            "view-projetos-cadastrar",
            "view-projetos-consultar",
            "view-projetos-dashboard",
            "view-projetos-setores",
            "view-cadastros-script-perguntas",
            "view-cadastros-cargos",
            "view-cadastros-usuarios",
            "view-cadastros-perfis",
        ]
        self._init_data_store()

        self.app = Flask(
            __name__,
            template_folder=str(Path(__file__).resolve().parent.parent / "web" / "templates"),
            static_folder=str(Path(__file__).resolve().parent.parent / "web" / "static"),
        )
        self.app.config["JSON_AS_ASCII"] = False
        self.app.config["SEND_FILE_MAX_AGE_DEFAULT"] = 0
        self.app.secret_key = os.getenv("APP_SECRET_KEY", "hpo-dev-secret-key")
        self.app.config["SESSION_COOKIE_HTTPONLY"] = True
        self.app.config["SESSION_COOKIE_SAMESITE"] = os.getenv("SESSION_COOKIE_SAMESITE", "Lax")
        self.app.config["SESSION_COOKIE_SECURE"] = self.is_production
        self.app.config["PERMANENT_SESSION_LIFETIME"] = timedelta(hours=int(os.getenv("SESSION_LIFETIME_HOURS", "12")))
        self._register_routes()

    def _register_routes(self):
        app = self.app

        @app.before_request
        def security_and_auth_guards():
            g.request_id = secrets.token_hex(8)
            session.permanent = True

            path = request.path or ""
            method = (request.method or "GET").upper()
            if "csrf_token" not in session:
                session["csrf_token"] = self._generate_csrf_token()

            if self._should_validate_csrf(path, method):
                sent_token = request.headers.get("X-CSRF-Token", "") or ""
                if not sent_token or sent_token != session.get("csrf_token", ""):
                    return jsonify({"ok": False, "message": "Falha de CSRF"}), 403

            if path.startswith("/static/") or path.startswith("/api/auth/") or path == "/login":
                return None
            user = self._get_current_user()
            if user:
                return None
            if path.startswith("/api/"):
                return jsonify({"ok": False, "message": "Nao autenticado"}), 401
            return redirect(url_for("login_page"))

        @app.after_request
        def add_security_headers(response):
            if not self.enable_security_headers:
                return response
            csp = (
                "default-src 'self'; "
                "script-src 'self' 'unsafe-inline' https://cdn.jsdelivr.net; "
                "style-src 'self' 'unsafe-inline'; "
                "img-src 'self' data:; "
                "font-src 'self' data:; "
                "connect-src 'self'; "
                "frame-ancestors 'none'; "
                "base-uri 'self'; "
                "form-action 'self'"
            )
            response.headers["Content-Security-Policy"] = csp
            response.headers["X-Frame-Options"] = "DENY"
            response.headers["X-Content-Type-Options"] = "nosniff"
            response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
            response.headers["Permissions-Policy"] = "camera=(), microphone=(), geolocation=()"
            response.headers["X-Request-ID"] = getattr(g, "request_id", "")
            if self.is_production:
                response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"
            # Evita servir HTML/JS/CSS antigos do cache durante evolucao local.
            content_type = str(response.headers.get("Content-Type", "")).lower()
            if any(x in content_type for x in ("text/html", "application/javascript", "text/css")):
                response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
                response.headers["Pragma"] = "no-cache"
                response.headers["Expires"] = "0"
            return response

        @app.get("/login")
        def login_page():
            return render_template("login.html", csrf_token=session.get("csrf_token", ""))

        @app.get("/")
        def index():
            return render_template("index.html", csrf_token=session.get("csrf_token", ""))

        @app.post("/api/auth/login")
        def auth_login():
            data = request.get_json(silent=True) or {}
            username = str(data.get("username", "")).strip().lower()
            password = str(data.get("password", ""))
            if not username or not password:
                return jsonify({"ok": False, "message": "Usuario e senha sao obrigatorios"}), 400

            users = self._load_json(self.users_file, default=[])
            user = next(
                (
                    u
                    for u in users
                    if str(u.get("username", "")).lower() == username
                    or str(u.get("email", "")).lower() == username
                ),
                None,
            )
            if not user or user.get("password_hash") != self._hash_password(password):
                return jsonify({"ok": False, "message": "Credenciais invalidas"}), 401

            if not bool(user.get("active", True)):
                return jsonify({"ok": False, "message": "Usuario inativo"}), 403
            if not bool(user.get("has_login", True)):
                return jsonify({"ok": False, "message": "Usuario sem acesso ao sistema"}), 403
            if user.get("must_change_password") and self._is_temp_password_expired(user):
                return jsonify({"ok": False, "message": "Senha temporaria expirada. Solicite novo reset."}), 403

            session["user_id"] = user["id"]
            session["csrf_token"] = self._generate_csrf_token()
            me = self._serialize_user_session(user)
            return jsonify({"ok": True, "user": me})

        @app.post("/api/auth/logout")
        def auth_logout():
            session.pop("user_id", None)
            session["csrf_token"] = self._generate_csrf_token()
            return jsonify({"ok": True})

        @app.get("/api/auth/me")
        def auth_me():
            user = self._get_current_user()
            if not user:
                return jsonify({"ok": False, "message": "Nao autenticado"}), 401
            return jsonify({"ok": True, "user": self._serialize_user_session(user), "available_views": self.available_views})

        @app.post("/api/auth/change-password")
        def auth_change_password():
            user = self._get_current_user()
            if not user:
                return jsonify({"ok": False, "message": "Nao autenticado"}), 401
            data = request.get_json(silent=True) or {}
            current_password = str(data.get("current_password", ""))
            new_password = str(data.get("new_password", "")).strip()
            if len(new_password) < 6:
                return jsonify({"ok": False, "message": "Nova senha deve ter ao menos 6 caracteres"}), 400

            if user.get("username") != "admin":
                if user.get("password_hash") != self._hash_password(current_password):
                    return jsonify({"ok": False, "message": "Senha atual invalida"}), 400

            users = self._load_json(self.users_file, default=[])
            idx = next((i for i, u in enumerate(users) if u.get("id") == user.get("id")), -1)
            if idx < 0:
                return jsonify({"ok": False, "message": "Usuario nao encontrado"}), 404
            users[idx]["password_hash"] = self._hash_password(new_password)
            users[idx]["must_change_password"] = False
            users[idx]["password_is_temporary"] = False
            users[idx]["temp_password_expires_at"] = ""
            users[idx]["updated_at"] = datetime.now().isoformat()
            self._save_json(self.users_file, users)
            return jsonify({"ok": True})

        @app.get("/api/roles")
        def list_roles():
            ok_resp = self._require_permission("view-cadastros-perfis")
            if ok_resp:
                return ok_resp
            roles = self._load_json(self.roles_file, default=[])
            return jsonify({"ok": True, "roles": roles, "available_views": self.available_views})

        @app.post("/api/roles")
        def create_role():
            ok_resp = self._require_permission("view-cadastros-perfis")
            if ok_resp:
                return ok_resp
            data = request.get_json(silent=True) or {}
            nome = str(data.get("nome", "")).strip()
            permissions = data.get("permissions") or []
            if not nome:
                return jsonify({"ok": False, "message": "Nome do papel obrigatorio"}), 400
            roles = self._load_json(self.roles_file, default=[])
            if any(str(r.get("nome", "")).lower() == nome.lower() for r in roles):
                return jsonify({"ok": False, "message": "Papel ja existe"}), 400
            role = {"id": str(uuid4()), "nome": nome, "permissions": [p for p in permissions if p in self.available_views]}
            roles.append(role)
            self._save_json(self.roles_file, roles)
            return jsonify({"ok": True, "role": role})

        @app.put("/api/roles/<role_id>")
        def update_role(role_id: str):
            ok_resp = self._require_permission("view-cadastros-perfis")
            if ok_resp:
                return ok_resp
            data = request.get_json(silent=True) or {}
            roles = self._load_json(self.roles_file, default=[])
            idx = next((i for i, r in enumerate(roles) if r.get("id") == role_id), -1)
            if idx < 0:
                return jsonify({"ok": False, "message": "Papel nao encontrado"}), 404
            if roles[idx].get("nome") == "Administrador":
                return jsonify({"ok": False, "message": "Papel Administrador nao pode ser alterado"}), 400
            nome = str(data.get("nome", roles[idx].get("nome", ""))).strip()
            permissions = data.get("permissions") or []
            roles[idx]["nome"] = nome or roles[idx].get("nome", "")
            roles[idx]["permissions"] = [p for p in permissions if p in self.available_views]
            self._save_json(self.roles_file, roles)
            return jsonify({"ok": True, "role": roles[idx]})

        @app.get("/api/users")
        def list_users():
            has_users_perm = self._require_permission("view-cadastros-usuarios") is None
            has_process_perm = self._require_permission("view-processos-cadastro") is None
            has_project_register_perm = self._require_permission("view-projetos-cadastrar") is None
            has_project_consult_perm = self._require_permission("view-projetos-consultar") is None
            has_project_dashboard_perm = self._require_permission("view-projetos-dashboard") is None
            if not has_users_perm and not has_process_perm and not has_project_register_perm and not has_project_consult_perm and not has_project_dashboard_perm:
                return jsonify({"ok": False, "message": "Sem permissao"}), 403
            login_access = str(request.args.get("login_access", "")).strip().lower()
            users = self._load_json(self.users_file, default=[])
            roles = self._load_json(self.roles_file, default=[])
            departamentos = self._load_json(self.sectors_file, default=[])
            cargos = self._load_json(self.cargos_file, default=[])
            role_map = {r.get("id"): r.get("nome") for r in roles}
            dept_map = {d.get("id"): d.get("nome") for d in departamentos}
            cargo_map = {c.get("id"): c.get("nome") for c in cargos}
            payload = [
                {
                    "id": u.get("id"),
                    "nome": u.get("nome"),
                    "email": u.get("email"),
                    "username": u.get("username"),
                    "role_id": u.get("role_id"),
                    "role_nome": role_map.get(u.get("role_id"), ""),
                    "departamento_id": u.get("departamento_id", ""),
                    "departamento_nome": dept_map.get(u.get("departamento_id"), ""),
                    "cargo_id": u.get("cargo_id", ""),
                    "cargo_nome": cargo_map.get(u.get("cargo_id"), ""),
                    "active": bool(u.get("active", True)),
                    "has_login": bool(u.get("has_login", True)),
                    "must_change_password": bool(u.get("must_change_password", False)),
                }
                for u in users
            ]
            if login_access in {"sim", "nao"}:
                want = login_access == "sim"
                payload = [u for u in payload if bool(u.get("has_login", True)) == want]
            return jsonify({"ok": True, "users": payload})

        @app.post("/api/users")
        def create_user():
            ok_resp = self._require_permission("view-cadastros-usuarios")
            if ok_resp:
                return ok_resp
            data = request.get_json(silent=True) or {}
            nome = str(data.get("nome", "")).strip()
            email = str(data.get("email", "")).strip().lower()
            role_id = str(data.get("role_id", "")).strip()
            has_login = bool(data.get("has_login", True))
            departamento_id = str(data.get("departamento_id", "")).strip()
            cargo_id = str(data.get("cargo_id", "")).strip()
            if not nome or not email:
                return jsonify({"ok": False, "message": "Nome e email sao obrigatorios"}), 400
            if has_login and not role_id:
                return jsonify({"ok": False, "message": "Papel e obrigatorio para usuarios com login"}), 400
            roles = self._load_json(self.roles_file, default=[])
            if has_login and not any(r.get("id") == role_id for r in roles):
                return jsonify({"ok": False, "message": "Papel invalido"}), 400
            departamentos = self._load_json(self.sectors_file, default=[])
            cargos = self._load_json(self.cargos_file, default=[])
            if departamento_id and not any(d.get("id") == departamento_id for d in departamentos):
                return jsonify({"ok": False, "message": "Departamento invalido"}), 400
            if cargo_id and not any(c.get("id") == cargo_id for c in cargos):
                return jsonify({"ok": False, "message": "Cargo invalido"}), 400
            users = self._load_json(self.users_file, default=[])
            if any(str(u.get("email", "")).lower() == email for u in users):
                return jsonify({"ok": False, "message": "Email ja cadastrado"}), 400
            username = ""
            temp_password = ""
            must_change_password = False
            password_hash = ""
            if has_login:
                username = email.split("@")[0]
                if any(str(u.get("username", "")).lower() == username for u in users):
                    username = f"{username}_{len(users)+1}"
                temp_password = self._generate_temporary_password()
                must_change_password = True
                password_hash = self._hash_password(temp_password)
            user = {
                "id": str(uuid4()),
                "nome": nome,
                "email": email,
                "username": username,
                "role_id": role_id,
                "departamento_id": departamento_id,
                "cargo_id": cargo_id,
                "has_login": has_login,
                "password_hash": password_hash,
                "must_change_password": must_change_password,
                "password_is_temporary": has_login,
                "temp_password_expires_at": (
                    datetime.now() + timedelta(hours=int(os.getenv("TEMP_PASSWORD_EXP_HOURS", "24")))
                ).isoformat() if has_login else "",
                "active": True,
                "created_at": datetime.now().isoformat(),
                "updated_at": datetime.now().isoformat(),
            }
            users.append(user)
            self._save_json(self.users_file, users)
            return jsonify(
                {
                    "ok": True,
                    "user": self._serialize_user_public(user),
                    "temp_password": temp_password,
                    "temp_password_expires_at": user.get("temp_password_expires_at", ""),
                }
            )

        @app.put("/api/users/<user_id>")
        def update_user(user_id: str):
            ok_resp = self._require_permission("view-cadastros-usuarios")
            if ok_resp:
                return ok_resp

            data = request.get_json(silent=True) or {}
            nome = str(data.get("nome", "")).strip()
            email = str(data.get("email", "")).strip().lower()
            role_id = str(data.get("role_id", "")).strip()
            has_login = bool(data.get("has_login", True))
            departamento_id = str(data.get("departamento_id", "")).strip()
            cargo_id = str(data.get("cargo_id", "")).strip()
            if not nome or not email:
                return jsonify({"ok": False, "message": "Nome e email sao obrigatorios"}), 400

            users = self._load_json(self.users_file, default=[])
            idx = next((i for i, u in enumerate(users) if u.get("id") == user_id), -1)
            if idx < 0:
                return jsonify({"ok": False, "message": "Usuario nao encontrado"}), 404

            if any(str(u.get("email", "")).lower() == email and str(u.get("id", "")) != user_id for u in users):
                return jsonify({"ok": False, "message": "Email ja cadastrado"}), 400

            roles = self._load_json(self.roles_file, default=[])
            departamentos = self._load_json(self.sectors_file, default=[])
            cargos = self._load_json(self.cargos_file, default=[])
            if has_login and not role_id:
                return jsonify({"ok": False, "message": "Papel e obrigatorio para usuarios com login"}), 400
            if role_id and not any(r.get("id") == role_id for r in roles):
                return jsonify({"ok": False, "message": "Papel invalido"}), 400
            if departamento_id and not any(d.get("id") == departamento_id for d in departamentos):
                return jsonify({"ok": False, "message": "Departamento invalido"}), 400
            if cargo_id and not any(c.get("id") == cargo_id for c in cargos):
                return jsonify({"ok": False, "message": "Cargo invalido"}), 400

            current = users[idx]
            current_user = self._get_current_user()
            if current_user and str(current_user.get("id", "")) == user_id and not has_login:
                return jsonify({"ok": False, "message": "Voce nao pode remover o proprio acesso ao sistema"}), 400
            if str(current.get("username", "")).strip().lower() == "admin" and not has_login:
                return jsonify({"ok": False, "message": "Usuario admin deve manter acesso ao sistema"}), 400

            temp_password = ""
            if has_login and not bool(current.get("has_login", True)):
                username = email.split("@")[0]
                if any(str(u.get("username", "")).lower() == username and str(u.get("id", "")) != user_id for u in users):
                    username = f"{username}_{len(users)+1}"
                temp_password = self._generate_temporary_password()
                current["username"] = username
                current["password_hash"] = self._hash_password(temp_password)
                current["must_change_password"] = True
                current["password_is_temporary"] = True
                current["temp_password_expires_at"] = (
                    datetime.now() + timedelta(hours=int(os.getenv("TEMP_PASSWORD_EXP_HOURS", "24")))
                ).isoformat()
            elif not has_login:
                current["username"] = ""
                current["role_id"] = ""
                current["password_hash"] = ""
                current["must_change_password"] = False
                current["password_is_temporary"] = False
                current["temp_password_expires_at"] = ""

            current["nome"] = nome
            current["email"] = email
            current["departamento_id"] = departamento_id
            current["cargo_id"] = cargo_id
            current["has_login"] = has_login
            current["role_id"] = role_id if has_login else ""
            current["updated_at"] = datetime.now().isoformat()

            self._save_json(self.users_file, users)
            return jsonify(
                {
                    "ok": True,
                    "user": self._serialize_user_public(current),
                    "temp_password": temp_password,
                    "temp_password_expires_at": current.get("temp_password_expires_at", ""),
                }
            )

        @app.put("/api/users/<user_id>/role")
        def update_user_role(user_id: str):
            ok_resp = self._require_permission("view-cadastros-usuarios")
            if ok_resp:
                return ok_resp
            data = request.get_json(silent=True) or {}
            role_id = str(data.get("role_id", "")).strip()
            roles = self._load_json(self.roles_file, default=[])
            if not any(r.get("id") == role_id for r in roles):
                return jsonify({"ok": False, "message": "Papel invalido"}), 400
            users = self._load_json(self.users_file, default=[])
            idx = next((i for i, u in enumerate(users) if u.get("id") == user_id), -1)
            if idx < 0:
                return jsonify({"ok": False, "message": "Usuario nao encontrado"}), 404
            if users[idx].get("username") == "admin":
                return jsonify({"ok": False, "message": "Usuario admin nao pode mudar de papel"}), 400
            users[idx]["role_id"] = role_id
            users[idx]["updated_at"] = datetime.now().isoformat()
            self._save_json(self.users_file, users)
            return jsonify({"ok": True})

        @app.post("/api/users/<user_id>/reset-password")
        def reset_user_password(user_id: str):
            ok_resp = self._require_permission("view-cadastros-usuarios")
            if ok_resp:
                return ok_resp
            users = self._load_json(self.users_file, default=[])
            idx = next((i for i, u in enumerate(users) if u.get("id") == user_id), -1)
            if idx < 0:
                return jsonify({"ok": False, "message": "Usuario nao encontrado"}), 404
            if not bool(users[idx].get("has_login", True)):
                return jsonify({"ok": False, "message": "Usuario sem login nao possui senha para reset"}), 400
            temp_password = self._generate_temporary_password()
            users[idx]["password_hash"] = self._hash_password(temp_password)
            users[idx]["must_change_password"] = users[idx].get("username") != "admin"
            users[idx]["password_is_temporary"] = True
            users[idx]["temp_password_expires_at"] = (
                datetime.now() + timedelta(hours=int(os.getenv("TEMP_PASSWORD_EXP_HOURS", "24")))
            ).isoformat()
            users[idx]["updated_at"] = datetime.now().isoformat()
            self._save_json(self.users_file, users)
            return jsonify(
                {
                    "ok": True,
                    "temp_password": temp_password,
                    "temp_password_expires_at": users[idx].get("temp_password_expires_at", ""),
                }
            )

        @app.delete("/api/users/<user_id>")
        def delete_user(user_id: str):
            ok_resp = self._require_permission("view-cadastros-usuarios")
            if ok_resp:
                return ok_resp

            current_user = self._get_current_user()
            users = self._load_json(self.users_file, default=[])
            idx = next((i for i, u in enumerate(users) if u.get("id") == user_id), -1)
            if idx < 0:
                return jsonify({"ok": False, "message": "Usuario nao encontrado"}), 404

            user = users[idx]
            user_nome = str(user.get("nome", "")).strip()
            if str(user.get("username", "")).strip().lower() == "admin":
                return jsonify({"ok": False, "message": "Usuario admin nao pode ser excluido"}), 400
            if current_user and str(current_user.get("id", "")) == user_id:
                return jsonify({"ok": False, "message": "Voce nao pode excluir o proprio usuario"}), 400

            processos = self._load_json(self.processes_file, default=[])
            processos_alterados = 0
            etapas_processo_alteradas = 0
            for processo in processos:
                processo_changed = False
                if str(processo.get("responsavel_id", "")) == user_id or str(processo.get("responsavel_nome", "")).strip() == user_nome:
                    processo["responsavel_id"] = ""
                    processo["responsavel_nome"] = ""
                    processo_changed = True
                for etapa in processo.get("etapas") or []:
                    if str(etapa.get("responsavel_id", "")) == user_id or str(etapa.get("responsavel_nome", "")).strip() == user_nome:
                        etapa["responsavel_id"] = ""
                        etapa["responsavel_nome"] = ""
                        etapas_processo_alteradas += 1
                        processo_changed = True
                if processo_changed:
                    processo["updated_at"] = datetime.now().isoformat()
                    processos_alterados += 1

            projects = self._load_json(self.projects_file, default=[])
            projetos_alterados = 0
            participantes_removidos = 0
            for project in projects:
                project_changed = False
                if str(project.get("responsavel_id", "")) == user_id or str(project.get("responsavel", "")).strip() == user_nome:
                    project["responsavel_id"] = ""
                    project["responsavel"] = ""
                    project_changed = True
                if str(project.get("focal_id", "")) == user_id or str(project.get("focal", "")).strip() == user_nome:
                    project["focal_id"] = ""
                    project["focal"] = ""
                    project_changed = True

                participantes_ids = [str(pid) for pid in (project.get("participantes_ids") or [])]
                if user_id in participantes_ids:
                    project["participantes_ids"] = [pid for pid in participantes_ids if pid != user_id]
                    participantes_removidos += 1
                    project_changed = True

                participantes = [str(nome).strip() for nome in (project.get("participantes") or []) if str(nome).strip()]
                if user_nome and user_nome in participantes:
                    project["participantes"] = [nome for nome in participantes if nome != user_nome]
                    project["setores_impactados"] = [nome for nome in (project.get("setores_impactados") or []) if str(nome).strip() != user_nome]
                    project_changed = True

                for etapa in project.get("etapas") or []:
                    if str(etapa.get("responsavel", "")).strip() == user_nome:
                        etapa["responsavel"] = ""
                        project_changed = True

                if project_changed:
                    project["updated_at"] = datetime.now().isoformat()
                    projetos_alterados += 1

            if processos_alterados or etapas_processo_alteradas:
                self._save_json(self.processes_file, processos)
            if projetos_alterados or participantes_removidos:
                self._save_json(self.projects_file, projects)

            users.pop(idx)
            self._save_json(self.users_file, users)
            return jsonify(
                {
                    "ok": True,
                    "cleanup": {
                        "processos_alterados": processos_alterados,
                        "etapas_processo_alteradas": etapas_processo_alteradas,
                        "projetos_alterados": projetos_alterados,
                        "participantes_removidos": participantes_removidos,
                    },
                }
            )

        @app.get("/api/state")
        def get_state():
            with self.lock:
                return jsonify(self._serialize_state())

        @app.get("/api/logs")
        def get_logs():
            return jsonify(
                {
                    "detailed": self.logger.get_log_history()[-400:],
                    "summary": self.logger.get_summary_history()[-200:],
                }
            )

        @app.post("/api/api-test")
        def test_api():
            ok, message = self.content_generator.test_api_connection()
            with self.lock:
                self.api_status = {"ok": ok, "message": message}
            return jsonify(self.api_status)

        @app.post("/api/priorizacao/calculate")
        def calculate_priorizacao():
            data = request.get_json(silent=True) or {}
            ok, payload = self._calculate_priorizacao_payload(data)
            if not ok:
                return jsonify({"ok": False, "message": payload}), 400
            self.logger.summary(
                f"Priorizacao calculada: score={payload['score_final']} prioridade={payload['prioridade']}"
            )
            return jsonify({"ok": True, "result": payload})

        @app.post("/api/priorizacao/export/word")
        def export_priorizacao_word():
            data = request.get_json(silent=True) or {}
            responsavel = str(data.get("responsavel", "")).strip()
            cargo = str(data.get("cargo", "")).strip()
            resumo_processo = str(data.get("resumo_processo", "")).strip()

            ok, result_or_error = self._calculate_priorizacao_payload(data)
            if not ok:
                return jsonify({"ok": False, "message": result_or_error}), 400
            result = result_or_error

            doc = Document()
            doc.add_heading("Relatorio de Priorizacao de Desenvolvimento", level=1)
            doc.add_paragraph(f"Responsavel: {responsavel or 'Nao informado'}")
            doc.add_paragraph(f"Cargo: {cargo or 'Nao informado'}")
            doc.add_paragraph(f"Resumo do processo: {resumo_processo or 'Nao informado'}")
            doc.add_paragraph("")
            doc.add_heading("Parametros de Entrada", level=2)
            doc.add_paragraph(f"Qtd Pessoas: {data.get('qtd_pessoas', 0)}")
            doc.add_paragraph(f"Horas Mensais: {data.get('horas_mensais', 0)}")
            doc.add_paragraph(f"Custo Mensal: {data.get('custo_mensal', 0)}")
            doc.add_paragraph(f"Custo Desenvolvimento: {data.get('custo_desenvolvimento', 0)}")
            doc.add_paragraph(f"Complexidade: {data.get('complexidade', 0)}")
            doc.add_paragraph(f"Desenvolvimento Interno: {data.get('dev_interno', 0)}")
            doc.add_paragraph("")
            doc.add_heading("Resultado do Calculo", level=2)
            doc.add_paragraph(f"Score Final: {result['score_final']}")
            doc.add_paragraph(f"Prioridade: {result['prioridade']}")
            doc.add_paragraph(f"Payback (meses): {result['payback_meses']}")
            doc.add_paragraph(f"CP Anual: {result['cp_anual']}")
            doc.add_paragraph(f"Economia Mensal: {result['economia_mensal']}")
            doc.add_paragraph(f"Custo Hora: {result['custo_hora']}")
            doc.add_paragraph(f"CP Score: {result['cp_score']}")
            doc.add_paragraph(f"Payback Score: {result['pb_score']}")
            doc.add_paragraph("")
            doc.add_heading("Regras de Priorizacao Aplicadas", level=2)
            doc.add_paragraph("Pesos: CP 35%, Payback 30%, Complexidade invertida 20%, Desenvolvimento Interno 15%.")
            doc.add_paragraph(
                "Score final = 0.35*cp_score + 0.30*pb_score + 0.20*(6-complexidade) + 0.15*dev_interno."
            )

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            self.logger.summary("Exportacao Word de priorizacao gerada")
            return send_file(
                buffer,
                as_attachment=True,
                download_name="Relatorio_Priorizacao.docx",
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        @app.get("/api/setores")
        def list_setores():
            setores = self._load_json(self.sectors_file, default=[])
            return jsonify({"ok": True, "setores": setores})

        @app.post("/api/setores")
        def create_setor():
            data = request.get_json(silent=True) or {}
            nome = str(data.get("nome", "")).strip()
            if not nome:
                return jsonify({"ok": False, "message": "Nome do setor é obrigatório"}), 400

            setores = self._load_json(self.sectors_file, default=[])
            if any(s["nome"].lower() == nome.lower() for s in setores):
                return jsonify({"ok": False, "message": "Setor já cadastrado"}), 400

            new_setor = {"id": str(uuid4()), "nome": nome}
            setores.append(new_setor)
            self._save_json(self.sectors_file, setores)
            self.logger.summary(f"Setor cadastrado: {nome}")
            return jsonify({"ok": True, "setor": new_setor})

        @app.delete("/api/setores/<setor_id>")
        def delete_setor(setor_id: str):
            setores = self._load_json(self.sectors_file, default=[])
            filtered = [s for s in setores if s.get("id") != setor_id]
            if len(filtered) == len(setores):
                return jsonify({"ok": False, "message": "Setor não encontrado"}), 404
            self._save_json(self.sectors_file, filtered)
            return jsonify({"ok": True})

        @app.get("/api/departamentos")
        def list_departamentos():
            has_cadastro_perm = (
                self._require_permission("view-projetos-setores") is None
            )
            has_processo_perm = self._require_permission("view-processos-cadastro") is None
            if not has_cadastro_perm and not has_processo_perm:
                return jsonify({"ok": False, "message": "Sem permissao"}), 403
            departamentos = self._load_json(self.sectors_file, default=[])
            return jsonify({"ok": True, "departamentos": departamentos})

        @app.post("/api/departamentos")
        def create_departamento():
            ok_resp = self._require_permission("view-projetos-setores")
            if ok_resp:
                return ok_resp
            data = request.get_json(silent=True) or {}
            nome = str(data.get("nome", "")).strip()
            if not nome:
                return jsonify({"ok": False, "message": "Nome do departamento e obrigatorio"}), 400
            departamentos = self._load_json(self.sectors_file, default=[])
            if any(str(d.get("nome", "")).lower() == nome.lower() for d in departamentos):
                return jsonify({"ok": False, "message": "Departamento ja cadastrado"}), 400
            departamento = {"id": str(uuid4()), "nome": nome}
            departamentos.append(departamento)
            self._save_json(self.sectors_file, departamentos)
            return jsonify({"ok": True, "departamento": departamento})

        @app.delete("/api/departamentos/<departamento_id>")
        def delete_departamento(departamento_id: str):
            ok_resp = self._require_permission("view-projetos-setores")
            if ok_resp:
                return ok_resp
            departamentos = self._load_json(self.sectors_file, default=[])
            filtered = [d for d in departamentos if d.get("id") != departamento_id]
            if len(filtered) == len(departamentos):
                return jsonify({"ok": False, "message": "Departamento nao encontrado"}), 404
            self._save_json(self.sectors_file, filtered)
            return jsonify({"ok": True})

        @app.get("/api/cargos")
        def list_cargos():
            has_cadastro_perm = self._require_permission("view-cadastros-cargos") is None
            has_processo_perm = self._require_permission("view-processos-cadastro") is None
            if not has_cadastro_perm and not has_processo_perm:
                return jsonify({"ok": False, "message": "Sem permissao"}), 403
            cargos = self._load_json(self.cargos_file, default=[])
            return jsonify({"ok": True, "cargos": cargos})

        @app.post("/api/cargos")
        def create_cargo():
            ok_resp = self._require_permission("view-cadastros-cargos")
            if ok_resp:
                return ok_resp
            data = request.get_json(silent=True) or {}
            nome = str(data.get("nome", "")).strip()
            if not nome:
                return jsonify({"ok": False, "message": "Nome do cargo e obrigatorio"}), 400
            cargos = self._load_json(self.cargos_file, default=[])
            if any(str(c.get("nome", "")).lower() == nome.lower() for c in cargos):
                return jsonify({"ok": False, "message": "Cargo ja cadastrado"}), 400
            cargo = {"id": str(uuid4()), "nome": nome}
            cargos.append(cargo)
            self._save_json(self.cargos_file, cargos)
            return jsonify({"ok": True, "cargo": cargo})

        @app.delete("/api/cargos/<cargo_id>")
        def delete_cargo(cargo_id: str):
            ok_resp = self._require_permission("view-cadastros-cargos")
            if ok_resp:
                return ok_resp
            cargos = self._load_json(self.cargos_file, default=[])
            filtered = [c for c in cargos if c.get("id") != cargo_id]
            if len(filtered) == len(cargos):
                return jsonify({"ok": False, "message": "Cargo nao encontrado"}), 404
            self._save_json(self.cargos_file, filtered)
            return jsonify({"ok": True})

        @app.get("/api/processos")
        def list_processos():
            has_cadastro_perm = self._require_permission("view-processos-cadastro") is None
            has_consulta_perm = self._require_permission("view-processos-consulta") is None
            if not has_cadastro_perm and not has_consulta_perm:
                return jsonify({"ok": False, "message": "Sem permissao"}), 403
            processos = self._load_json(self.processes_file, default=[])
            nome = str(request.args.get("nome", "")).strip().lower()
            departamento_id = str(request.args.get("departamento_id", "")).strip()
            status = str(request.args.get("status", "")).strip().lower()
            filtered = []
            for p in processos:
                if nome and nome not in str(p.get("nome", "")).lower():
                    continue
                if departamento_id and str(p.get("departamento_id", "")) != departamento_id:
                    continue
                if status and str(p.get("status", "")).lower() != status:
                    continue
                filtered.append(p)
            return jsonify({"ok": True, "processos": filtered})

        @app.post("/api/processos")
        def create_processo():
            ok_resp = self._require_permission("view-processos-cadastro")
            if ok_resp:
                return ok_resp
            data = request.get_json(silent=True) or {}
            ok, payload_or_error = self._validate_processo_payload(data, None)
            if not ok:
                return jsonify({"ok": False, "message": payload_or_error}), 400
            payload = payload_or_error
            payload["id"] = str(uuid4())
            for etapa in payload.get("etapas", []):
                etapa["processo_id"] = payload["id"]
            payload["data_criacao"] = datetime.now().isoformat()
            payload["ultima_atualizacao"] = datetime.now().isoformat()
            processos = self._load_json(self.processes_file, default=[])
            processos.append(payload)
            self._save_json(self.processes_file, processos)
            return jsonify({"ok": True, "processo": payload})

        @app.get("/api/processos/<processo_id>")
        def get_processo(processo_id: str):
            has_cadastro_perm = self._require_permission("view-processos-cadastro") is None
            has_consulta_perm = self._require_permission("view-processos-consulta") is None
            if not has_cadastro_perm and not has_consulta_perm:
                return jsonify({"ok": False, "message": "Sem permissao"}), 403
            processos = self._load_json(self.processes_file, default=[])
            processo = next((p for p in processos if p.get("id") == processo_id), None)
            if not processo:
                return jsonify({"ok": False, "message": "Processo nao encontrado"}), 404
            return jsonify({"ok": True, "processo": processo})

        @app.put("/api/processos/<processo_id>")
        def update_processo(processo_id: str):
            ok_resp = self._require_permission("view-processos-cadastro")
            if ok_resp:
                return ok_resp
            data = request.get_json(silent=True) or {}
            processos = self._load_json(self.processes_file, default=[])
            idx = next((i for i, p in enumerate(processos) if p.get("id") == processo_id), -1)
            if idx < 0:
                return jsonify({"ok": False, "message": "Processo nao encontrado"}), 404
            ok, payload_or_error = self._validate_processo_payload(data, processo_id)
            if not ok:
                return jsonify({"ok": False, "message": payload_or_error}), 400
            payload = payload_or_error
            payload["id"] = processo_id
            for etapa in payload.get("etapas", []):
                etapa["processo_id"] = processo_id
            payload["data_criacao"] = processos[idx].get("data_criacao") or datetime.now().isoformat()
            payload["ultima_atualizacao"] = datetime.now().isoformat()
            processos[idx] = payload
            self._save_json(self.processes_file, processos)
            return jsonify({"ok": True, "processo": payload})

        @app.get("/api/processos/dashboard")
        def processos_dashboard():
            has_cadastro_perm = self._require_permission("view-processos-cadastro") is None
            has_consulta_perm = self._require_permission("view-processos-consulta") is None
            has_dashboard_perm = self._require_permission("view-processos-dashboard") is None
            if not has_cadastro_perm and not has_consulta_perm and not has_dashboard_perm:
                return jsonify({"ok": False, "message": "Sem permissao"}), 403

            processos = self._load_json(self.processes_file, default=[])
            dept_filter = str(request.args.get("departamento_id", "")).strip()
            resp_filter = str(request.args.get("responsavel_id", "")).strip()
            status_filter = str(request.args.get("status", "")).strip().lower()

            filtered = []
            for p in processos:
                if dept_filter and str(p.get("departamento_id", "")) != dept_filter:
                    continue
                if resp_filter and str(p.get("responsavel_id", "")) != resp_filter:
                    continue
                if status_filter and str(p.get("status", "")).lower() != status_filter:
                    continue
                filtered.append(p)

            total_processos = len(filtered)
            total_etapas = sum(len(p.get("etapas") or []) for p in filtered)
            media_etapas = round((total_etapas / total_processos), 2) if total_processos else 0
            ativos = sum(1 for p in filtered if p.get("status") == "ativo")
            rascunhos = sum(1 for p in filtered if p.get("status") == "rascunho")
            inativos = sum(1 for p in filtered if p.get("status") == "inativo")

            departamentos = {}
            etapas_departamentos = {}
            responsaveis = {}
            status_map = {"ativo": 0, "rascunho": 0, "inativo": 0}

            for p in filtered:
                dept = str(p.get("departamento_nome", "")).strip() or "Não informado"
                resp = str(p.get("responsavel_nome", "")).strip() or "Não informado"
                departamentos[dept] = departamentos.get(dept, 0) + 1
                responsaveis[resp] = responsaveis.get(resp, 0) + 1
                etapas_departamentos[dept] = etapas_departamentos.get(dept, 0) + len(p.get("etapas") or [])
                st = str(p.get("status", "")).lower().strip()
                if st in status_map:
                    status_map[st] += 1

            return jsonify(
                {
                    "ok": True,
                    "kpis": {
                        "total_processos": total_processos,
                        "total_etapas": total_etapas,
                        "media_etapas": media_etapas,
                        "ativos": ativos,
                        "rascunhos": rascunhos,
                        "inativos": inativos,
                        "departamentos_com_processo": len(departamentos),
                        "responsaveis_ativos": len(responsaveis),
                    },
                    "charts": {
                        "processos_por_departamento": departamentos,
                        "etapas_por_departamento": etapas_departamentos,
                        "processos_por_responsavel": responsaveis,
                        "status_processos": status_map,
                    },
                    "processos": filtered,
                }
            )

        @app.post("/api/processos/<processo_id>/inativar")
        def inativar_processo(processo_id: str):
            ok_resp = self._require_permission("view-processos-cadastro")
            if ok_resp:
                return ok_resp
            processos = self._load_json(self.processes_file, default=[])
            idx = next((i for i, p in enumerate(processos) if p.get("id") == processo_id), -1)
            if idx < 0:
                return jsonify({"ok": False, "message": "Processo nao encontrado"}), 404
            processos[idx]["status"] = "inativo"
            processos[idx]["ultima_atualizacao"] = datetime.now().isoformat()
            self._save_json(self.processes_file, processos)
            return jsonify({"ok": True, "processo": processos[idx]})

        @app.post("/api/processos/<processo_id>/etapas/<etapa_id>/comentarios")
        def add_etapa_comment(processo_id: str, etapa_id: str):
            ok_resp = self._require_permission("view-processos-cadastro")
            if ok_resp:
                return ok_resp
            data = request.get_json(silent=True) or {}
            texto = str(data.get("texto", "")).strip()
            if not texto:
                return jsonify({"ok": False, "message": "Comentario obrigatorio"}), 400
            user = self._get_current_user()
            processos = self._load_json(self.processes_file, default=[])
            proc_idx = next((i for i, p in enumerate(processos) if p.get("id") == processo_id), -1)
            if proc_idx < 0:
                return jsonify({"ok": False, "message": "Processo nao encontrado"}), 404
            etapas = processos[proc_idx].get("etapas") or []
            etapa_idx = next((i for i, e in enumerate(etapas) if e.get("id") == etapa_id), -1)
            if etapa_idx < 0:
                return jsonify({"ok": False, "message": "Etapa nao encontrada"}), 404
            conhecimento = etapas[etapa_idx].setdefault("conhecimento", {})
            comentarios = conhecimento.setdefault("comentarios", [])
            comentario = {
                "id": str(uuid4()),
                "autor_id": user.get("id") if user else "",
                "autor_nome": (user or {}).get("nome", "Usuario"),
                "data_hora": datetime.now().isoformat(),
                "texto": texto,
                "edit_history": [],
            }
            comentarios.append(comentario)
            comentarios.sort(key=lambda c: str(c.get("data_hora", "")))
            processos[proc_idx]["ultima_atualizacao"] = datetime.now().isoformat()
            self._save_json(self.processes_file, processos)
            return jsonify({"ok": True, "comentario": comentario})

        @app.put("/api/processos/<processo_id>/etapas/<etapa_id>/comentarios/<comentario_id>")
        def edit_etapa_comment(processo_id: str, etapa_id: str, comentario_id: str):
            ok_resp = self._require_permission("view-processos-cadastro")
            if ok_resp:
                return ok_resp
            data = request.get_json(silent=True) or {}
            novo_texto = str(data.get("texto", "")).strip()
            if not novo_texto:
                return jsonify({"ok": False, "message": "Comentario obrigatorio"}), 400
            user = self._get_current_user()
            processos = self._load_json(self.processes_file, default=[])
            proc_idx = next((i for i, p in enumerate(processos) if p.get("id") == processo_id), -1)
            if proc_idx < 0:
                return jsonify({"ok": False, "message": "Processo nao encontrado"}), 404
            etapas = processos[proc_idx].get("etapas") or []
            etapa_idx = next((i for i, e in enumerate(etapas) if e.get("id") == etapa_id), -1)
            if etapa_idx < 0:
                return jsonify({"ok": False, "message": "Etapa nao encontrada"}), 404
            comentarios = ((etapas[etapa_idx].get("conhecimento") or {}).get("comentarios")) or []
            com_idx = next((i for i, c in enumerate(comentarios) if c.get("id") == comentario_id), -1)
            if com_idx < 0:
                return jsonify({"ok": False, "message": "Comentario nao encontrado"}), 404
            old = comentarios[com_idx]
            history = old.get("edit_history") or []
            history.append(
                {
                    "texto_anterior": old.get("texto", ""),
                    "editado_em": datetime.now().isoformat(),
                    "editado_por": (user or {}).get("nome", "Usuario"),
                }
            )
            old["texto"] = novo_texto
            old["edit_history"] = history
            processos[proc_idx]["ultima_atualizacao"] = datetime.now().isoformat()
            self._save_json(self.processes_file, processos)
            return jsonify({"ok": True, "comentario": old})

        @app.post("/api/projects")
        def create_project():
            data = request.get_json(silent=True) or {}
            ok, payload = self._validate_project_payload(data)
            if not ok:
                return jsonify({"ok": False, "message": payload}), 400

            projects = self._load_json(self.projects_file, default=[])
            payload["id"] = str(uuid4())
            payload["created_at"] = datetime.now().isoformat()
            payload["updated_at"] = datetime.now().isoformat()
            projects.append(payload)
            self._save_json(self.projects_file, projects)
            self.logger.summary(f"Projeto cadastrado: {payload['nome']}")
            return jsonify({"ok": True, "project": payload})

        @app.get("/api/projects")
        def list_projects():
            projects = self._load_json(self.projects_file, default=[])
            current_user = self._get_current_user()
            projects = self._filter_projects_for_user(projects, current_user)
            nome = str(request.args.get("nome", "")).strip().lower()
            status = str(request.args.get("status", "")).strip().lower()
            responsavel = str(request.args.get("responsavel", "")).strip().lower()
            data_ini = str(request.args.get("data_inicio", "")).strip()
            data_fim = str(request.args.get("data_fim", "")).strip()

            def in_range(proj: Dict) -> bool:
                dt = proj.get("data_inicio_previsto") or proj.get("dt_inicio_real") or ""
                if not dt:
                    return True
                if data_ini and dt < data_ini:
                    return False
                if data_fim and dt > data_fim:
                    return False
                return True

            filtered = []
            for p in projects:
                if nome and nome not in p.get("nome", "").lower():
                    continue
                if status and status != p.get("status", "").lower():
                    continue
                if responsavel and responsavel not in p.get("responsavel", "").lower():
                    continue
                if not in_range(p):
                    continue
                filtered.append(p)

            return jsonify({"ok": True, "projects": filtered})

        @app.get("/api/projects/<project_id>")
        def get_project(project_id: str):
            projects = self._load_json(self.projects_file, default=[])
            current_user = self._get_current_user()
            projects = self._filter_projects_for_user(projects, current_user)
            project = next((p for p in projects if p.get("id") == project_id), None)
            if not project:
                return jsonify({"ok": False, "message": "Projeto não encontrado"}), 404
            return jsonify({"ok": True, "project": project})

        @app.put("/api/projects/<project_id>")
        def update_project(project_id: str):
            data = request.get_json(silent=True) or {}
            action = str(data.get("_action", "")).strip().lower()
            if action == "delete":
                projects = self._load_json(self.projects_file, default=[])
                idx = next((i for i, p in enumerate(projects) if p.get("id") == project_id), -1)
                if idx < 0:
                    return jsonify({"ok": False, "message": "Projeto não encontrado"}), 404
                deleted = projects.pop(idx)
                self._save_json(self.projects_file, projects)
                self.logger.summary(f"Projeto excluído: {deleted.get('nome', project_id)}")
                return jsonify({"ok": True})
            ok, payload = self._validate_project_payload(data)
            if not ok:
                return jsonify({"ok": False, "message": payload}), 400

            projects = self._load_json(self.projects_file, default=[])
            idx = next((i for i, p in enumerate(projects) if p.get("id") == project_id), -1)
            if idx < 0:
                return jsonify({"ok": False, "message": "Projeto não encontrado"}), 404

            created_at = projects[idx].get("created_at")
            payload["id"] = project_id
            payload["created_at"] = created_at or datetime.now().isoformat()
            payload["updated_at"] = datetime.now().isoformat()
            projects[idx] = payload
            self._save_json(self.projects_file, projects)
            self.logger.summary(f"Projeto atualizado: {payload['nome']}")
            return jsonify({"ok": True, "project": payload})

        @app.post("/api/projects/<project_id>")
        def post_project_action(project_id: str):
            data = request.get_json(silent=True) or {}
            action = str(
                data.get("_action")
                or data.get("action")
                or data.get("_method")
                or request.args.get("action", "")
            ).strip().lower()
            if action in {"delete", "remove", "excluir"}:
                return delete_project(project_id)
            return update_project(project_id)

        @app.delete("/api/projects/<project_id>")
        def delete_project(project_id: str):
            projects = self._load_json(self.projects_file, default=[])
            idx = next((i for i, p in enumerate(projects) if p.get("id") == project_id), -1)
            if idx < 0:
                return jsonify({"ok": False, "message": "Projeto não encontrado"}), 404
            deleted = projects.pop(idx)
            self._save_json(self.projects_file, projects)
            self.logger.summary(f"Projeto excluído: {deleted.get('nome', project_id)}")
            return jsonify({"ok": True})

        @app.post("/api/projects/<project_id>/delete")
        def delete_project_post(project_id: str):
            return delete_project(project_id)

        @app.post("/api/projects/delete")
        def delete_project_by_payload():
            data = request.get_json(silent=True) or {}
            project_id = str(data.get("project_id", "")).strip()
            if not project_id:
                return jsonify({"ok": False, "message": "project_id obrigatorio"}), 400
            return delete_project(project_id)

        @app.delete("/api/projects/<project_id>/stages/<stage_id>")
        def delete_project_stage(project_id: str, stage_id: str):
            projects = self._load_json(self.projects_file, default=[])
            idx = next((i for i, p in enumerate(projects) if p.get("id") == project_id), -1)
            if idx < 0:
                return jsonify({"ok": False, "message": "Projeto não encontrado"}), 404

            project = projects[idx]
            stages = list(project.get("etapas") or [])
            stage_idx = next((i for i, s in enumerate(stages) if s.get("id") == stage_id), -1)
            if stage_idx < 0:
                return jsonify({"ok": False, "message": "Etapa não encontrada"}), 404

            stages.pop(stage_idx)
            project["etapas"] = stages
            project["progresso"] = self._compute_progress(stages, project.get("status"))
            project["tarefas"] = len(stages)
            project["tarefasConcluidas"] = sum(1 for s in stages if s.get("status") == "concluida")
            project["updated_at"] = datetime.now().isoformat()
            projects[idx] = project
            self._save_json(self.projects_file, projects)
            return jsonify({"ok": True, "project": project})

        @app.post("/api/projects/<project_id>/stages/<stage_id>")
        def project_stage_action(project_id: str, stage_id: str):
            data = request.get_json(silent=True) or {}
            action = str(
                data.get("_action")
                or data.get("action")
                or data.get("_method")
                or request.args.get("action", "")
            ).strip().lower()
            if action in {"delete", "remove", "excluir"}:
                return delete_project_stage(project_id, stage_id)
            return jsonify({"ok": False, "message": "Acao invalida para etapa"}), 400

        @app.post("/api/projects/<project_id>/stages/<stage_id>/delete")
        def delete_project_stage_post(project_id: str, stage_id: str):
            return delete_project_stage(project_id, stage_id)

        @app.post("/api/projects/stages/delete")
        def delete_project_stage_by_payload():
            data = request.get_json(silent=True) or {}
            project_id = str(data.get("project_id", "")).strip()
            stage_id = str(data.get("stage_id", "")).strip()
            if not project_id or not stage_id:
                return jsonify({"ok": False, "message": "project_id e stage_id obrigatorios"}), 400
            return delete_project_stage(project_id, stage_id)

        @app.get("/api/projects/dashboard")
        def projects_dashboard():
            projects = self._load_json(self.projects_file, default=[])
            current_user = self._get_current_user()
            projects = self._filter_projects_for_user(projects, current_user)
            total = len(projects)
            ativos = sum(1 for p in projects if p.get("status") in {"em_andamento", "proxima_sprint"})
            concluidos = sum(1 for p in projects if p.get("status") == "concluido")
            pausados = sum(1 for p in projects if p.get("status") == "pausado")
            criticos = sum(1 for p in projects if p.get("criticidade") in {"urgente", "critica"})
            atrasados = 0
            hoje = datetime.now().date().isoformat()
            for p in projects:
                prev = p.get("previsao_termino") or ""
                term = p.get("termino_real") or ""
                if prev and prev < hoje and not term and p.get("status") != "concluido":
                    atrasados += 1

            progresso_list = [float(p.get("progresso", 0) or 0) for p in projects]
            progresso_medio = round(sum(progresso_list) / len(progresso_list), 2) if progresso_list else 0
            orc_total = round(sum(float(p.get("orcamento", 0) or 0) for p in projects), 2)
            custo_total = round(sum(float(p.get("custo_atual", 0) or 0) for p in projects), 2)
            taxa_conclusao = round((concluidos / total) * 100, 2) if total else 0
            responsaveis_ativos = len({p.get("responsavel") for p in projects if p.get("responsavel")})
            setores_env = len({p.get("setor_projeto") for p in projects if p.get("setor_projeto")})

            return jsonify(
                {
                    "ok": True,
                    "kpis": {
                        "total_projetos": total,
                        "ativos": ativos,
                        "concluidos": concluidos,
                        "atrasados": atrasados,
                        "pausados": pausados,
                        "criticos": criticos,
                        "taxa_conclusao": taxa_conclusao,
                        "progresso_medio": progresso_medio,
                        "orcamento_total": orc_total,
                        "custo_total": custo_total,
                        "responsaveis_ativos": responsaveis_ativos,
                        "setores_envolvidos": setores_env,
                    },
                    "projects": projects,
                }
            )

        @app.post("/api/config")
        def save_config():
            data = request.get_json(silent=True) or {}
            origem = str(data.get("origem_path", "")).strip()
            destino = str(data.get("destino_path", "")).strip()

            if not origem or not destino:
                return jsonify({"ok": False, "message": "Origem e destino sao obrigatorios"}), 400

            if not Path(origem).exists():
                return jsonify({"ok": False, "message": "Pasta de origem nao existe"}), 400

            if not Path(destino).exists():
                return jsonify({"ok": False, "message": "Pasta de destino nao existe"}), 400

            with self.lock:
                self.origem_path = origem
                self.destino_path = destino

            self.logger.info("Configuracao atualizada de caminhos")
            self.logger.summary("Configuracao de caminhos atualizada")
            return jsonify({"ok": True})

        @app.post("/api/scan")
        def scan_documents():
            with self.lock:
                if not self.origem_path or not self.destino_path:
                    return jsonify({"ok": False, "message": "Configure origem e destino antes de escanear"}), 400
                origem = self.origem_path
                destino = self.destino_path

            docs = self.file_reader.scan_directory_structure(origem)
            self.logger.summary(f"Escaneamento concluido: {len(docs)} documentos encontrados")

            normalized_docs = []
            for doc in docs:
                output_path = Path(destino) / self.settings.get_output_folder_structure(
                    doc["departamento"], doc["subarea"], doc["filename"]
                )
                exists, existing_files = self.file_reader.check_output_files_exist(str(output_path))
                normalized_docs.append(
                    {
                        **doc,
                        "status": "ja_existe" if exists else "pendente",
                        "existing_files": existing_files,
                        "processed": exists,
                        "selected": not exists,
                        "preview_ready": False,
                    }
                )

            with self.lock:
                self.documents = normalized_docs
                self.processing = False
                self.current_index = 0
                self.current_preview = None
                self.pending_previews = {}

            return jsonify({"ok": True, "count": len(normalized_docs)})

        @app.post("/api/documents/select")
        def select_documents():
            data = request.get_json(silent=True) or {}
            mode = str(data.get("mode", "")).strip().lower()
            indices = data.get("indices", [])

            with self.lock:
                if mode == "all":
                    for doc in self.documents:
                        if not doc.get("processed"):
                            doc["selected"] = True
                elif mode == "none":
                    for doc in self.documents:
                        doc["selected"] = False
                elif mode == "invert":
                    for doc in self.documents:
                        if not doc.get("processed"):
                            doc["selected"] = not bool(doc.get("selected"))
                elif mode == "set":
                    selected_set = {int(i) for i in indices}
                    for idx, doc in enumerate(self.documents):
                        if doc.get("processed"):
                            continue
                        doc["selected"] = idx in selected_set
                else:
                    return jsonify({"ok": False, "message": "Modo de selecao invalido"}), 400

            return jsonify({"ok": True, "state": self._serialize_state()})

        @app.post("/api/process/start")
        def start_process():
            with self.lock:
                if not self.documents:
                    return jsonify({"ok": False, "message": "Nenhum documento escaneado"}), 400
                if self.processing:
                    return jsonify({"ok": False, "message": "Processamento ja esta em andamento"}), 400
                selected_pending = [
                    doc for doc in self.documents
                    if bool(doc.get("selected")) and not bool(doc.get("processed"))
                ]
                if not selected_pending:
                    return jsonify({"ok": False, "message": "Nenhum documento selecionado para processar"}), 400
                self.processing = True
                self.current_preview = None

            self.logger.summary("Processamento iniciado para documentos selecionados")
            self.worker_thread = threading.Thread(target=self._process_documents_async, daemon=True)
            self.worker_thread.start()
            return jsonify({"ok": True, "message": "Processamento iniciado"})

        @app.post("/api/process/stop")
        def stop_process():
            with self.lock:
                self.processing = False
            self.logger.info("Processamento interrompido pelo usuario")
            self.logger.summary("Processamento interrompido pelo usuario")
            return jsonify({"ok": True})

        @app.get("/api/process/preview/<int:index>")
        def get_preview(index: int):
            with self.lock:
                preview = self.pending_previews.get(index)
                state = self._serialize_state()
            if not preview:
                return jsonify({"ok": False, "message": "Preview nao encontrado para este documento"}), 404
            return jsonify({"ok": True, "preview": preview, "state": state})

        @app.post("/api/process/approve")
        def approve_preview():
            data = request.get_json(silent=True) or {}
            edited_docs = data.get("documents") or {}
            index = int(data.get("index", -1))

            with self.lock:
                preview = self.pending_previews.get(index)
                destino = self.destino_path

            if not preview:
                return jsonify({"ok": False, "message": "Nao ha preview ativo"}), 400

            required = {"ficha_tecnica", "fluxograma", "riscos"}
            if not required.issubset(set(edited_docs.keys())):
                return jsonify({"ok": False, "message": "Documentos incompletos para aprovacao"}), 400

            processo_name = preview["doc"]["filename"]
            output_path = Path(destino) / self.settings.get_output_folder_structure(
                preview["doc"]["departamento"], preview["doc"]["subarea"], processo_name
            )

            results = self.document_generator.create_all_documents(
                edited_docs, str(output_path), processo_name
            )
            success_count = sum(1 for status in results.values() if status)

            with self.lock:
                idx = preview["index"]
                if 0 <= idx < len(self.documents):
                    self.documents[idx]["processed"] = True
                    self.documents[idx]["status"] = "processado"
                    self.documents[idx]["preview_ready"] = False
                self.pending_previews.pop(idx, None)
                self.current_preview = None

            self.logger.info(f"{success_count}/3 documentos salvos para {processo_name}")
            self.logger.summary(f"Concluido: {processo_name} ({success_count}/3 arquivos salvos)")
            with self.lock:
                state = self._serialize_state()
            return jsonify({"ok": True, "state": state})

        @app.post("/api/process/cancel")
        def cancel_preview():
            with self.lock:
                preview = self.current_preview
                self.current_preview = None

            if preview:
                self.logger.info(f"Geracao cancelada para {preview['doc']['filename']}")
                self.logger.summary(f"Cancelado: {preview['doc']['filename']}")

            with self.lock:
                state = self._serialize_state()
            return jsonify({"ok": True, "state": state})

        @app.post("/api/process/regenerate")
        def regenerate_preview():
            data = request.get_json(silent=True) or {}
            index = int(data.get("index", -1))
            with self.lock:
                has_doc = 0 <= index < len(self.documents)
            if not has_doc:
                return jsonify({"ok": False, "message": "Indice de documento invalido"}), 400

            payload = self._create_preview_for_index(index)
            if not payload:
                return jsonify({"ok": False, "message": "Falha ao gerar preview"}), 500
            with self.lock:
                self.pending_previews[index] = payload
                self.documents[index]["status"] = "aguardando_aprovacao"
                self.documents[index]["preview_ready"] = True
                state = self._serialize_state()
            return jsonify({"ok": True, "preview": payload, "state": state})

        @app.post("/api/assistant/chat")
        def assistant_chat():
            user = self._get_current_user()
            if not user:
                return jsonify({"ok": False, "message": "Nao autenticado"}), 401

            data = request.get_json(silent=True) or {}
            message = str(data.get("message", "")).strip()
            if not message:
                return jsonify({"ok": False, "message": "Mensagem obrigatoria"}), 400

            permissions = self._get_user_permissions(user)
            context_payload = self._build_assistant_context_for_user(user, permissions)

            system_prompt = (
                "Voce e Justine, assistente virtual do HPO (Hub de Processos Operacionais). "
                "Responda em portugues, linguagem natural, objetiva, respeitosa e amigavel. "
                "Escreva de forma organizada, sem markdown pesado, com quebra de linha quando ajudar leitura. "
                "Use somente o contexto permitido abaixo. "
                "Se o usuario pedir dado sem permissao, informe claramente que ele nao tem acesso para essa consulta. "
                "Nao invente numeros, nomes ou status. "
                "Priorize leitura de dados: traga valores, totais e detalhamento por status/departamento/responsavel quando fizer sentido. "
                "Quando a pergunta envolver comentarios/anotacoes, retorne itens com origem, data, autor e texto."
            )

            context_text = (
                f"Usuario atual: {context_payload['user_name']}\n"
                f"Permissoes: {', '.join(context_payload['permissions']) if context_payload['permissions'] else 'nenhuma'}\n"
                f"Acesso dados projetos: {'sim' if context_payload['can_access_projects'] else 'nao'}\n"
                f"Acesso dados processos: {'sim' if context_payload['can_access_processes'] else 'nao'}\n"
                f"Acesso dados cadastros: {'sim' if context_payload['can_access_registers'] else 'nao'}\n\n"
                f"Resumo permitido:\n{context_payload['context_text']}\n"
            )

            try:
                response = self.content_generator.client.chat.completions.create(
                    model=self.settings.OPENAI_MODEL,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "system", "content": context_text},
                        {"role": "user", "content": message},
                    ],
                    max_tokens=500,
                    temperature=0.2,
                    timeout=self.settings.OPENAI_TIMEOUT,
                )
                answer = (response.choices[0].message.content or "").strip()
                if not answer:
                    answer = "Nao consegui gerar resposta agora. Tente novamente em instantes."
                return jsonify({"ok": True, "answer": answer})
            except Exception as exc:
                self.logger.warning(f"Falha no assistente OpenAI (fallback local): {exc}")
                fallback = self._assistant_local_fallback_answer(message, context_payload)
                return jsonify({"ok": True, "answer": fallback})

    def _serialize_state(self) -> Dict:
        total = len(self.documents)
        processed = sum(1 for item in self.documents if item.get("processed"))
        selected_total = sum(1 for item in self.documents if bool(item.get("selected")))
        selected_processed = sum(
            1 for item in self.documents if bool(item.get("selected")) and bool(item.get("processed"))
        )
        selected_pending = max(selected_total - selected_processed, 0)
        return {
            "origem_path": self.origem_path,
            "destino_path": self.destino_path,
            "processing": self.processing,
            "total_documents": total,
            "processed_documents": processed,
            "progress_percent": round((processed / total) * 100, 2) if total else 0,
            "selected_total": selected_total,
            "selected_processed": selected_processed,
            "selected_pending": selected_pending,
            "selected_progress_percent": round((selected_processed / selected_total) * 100, 2) if selected_total else 0,
            "has_preview": bool(self.pending_previews),
            "preview_doc": next(
                (self.documents[i]["filename"] for i in self.pending_previews if 0 <= i < len(self.documents)),
                "",
            ),
            "documents": self.documents,
            "api_status": self.api_status,
        }

    def _create_preview_for_index(self, index: int) -> Optional[Dict]:
        with self.lock:
            if index < 0 or index >= len(self.documents):
                return None
            doc = self.documents[index]

        processo_name = doc["filename"]
        self.logger.info(f"Processando: {processo_name}")
        self.logger.summary(f"Em processamento: {processo_name}")

        titulo, texto_completo, texto_normalizado = self.file_reader.read_document_content(doc["file_path"])
        if not texto_normalizado:
            self.logger.error(f"Falha ao ler {processo_name}")
            self.logger.summary(f"Falha de leitura: {processo_name}")
            return None

        chunks = self.chunking_engine.create_semantic_chunks(texto_normalizado)
        generated = self.content_generator.generate_all_documents(
            chunks, processo_name, doc["departamento"], doc["subarea"]
        )

        preview_payload = {
            "index": index,
            "doc": {
                "filename": doc["filename"],
                "departamento": doc["departamento"],
                "subarea": doc["subarea"],
            },
            "documents": generated,
            "metadata": {
                "title": titulo,
                "raw_text_size": len(texto_completo),
                "normalized_text_size": len(texto_normalizado),
                "chunks": len(chunks),
            },
        }
        return preview_payload

    def _process_documents_async(self):
        while True:
            with self.lock:
                if not self.processing:
                    break
                next_index = None
                for i, doc in enumerate(self.documents):
                    if bool(doc.get("selected")) and not bool(doc.get("processed")) and not bool(doc.get("preview_ready")):
                        next_index = i
                        self.documents[i]["status"] = "gerando"
                        break

            if next_index is None:
                with self.lock:
                    self.processing = False
                self.logger.summary("Geracao concluida para todos os selecionados")
                break

            payload = self._create_preview_for_index(next_index)

            with self.lock:
                if not self.processing:
                    break
                if payload:
                    self.pending_previews[next_index] = payload
                    self.documents[next_index]["status"] = "aguardando_aprovacao"
                    self.documents[next_index]["preview_ready"] = True
                    self.logger.summary(f"Preview pronto: {self.documents[next_index]['filename']}")
                else:
                    self.documents[next_index]["status"] = "erro_geracao"
                    self.documents[next_index]["preview_ready"] = False
                    self.logger.summary(f"Falha na geracao: {self.documents[next_index]['filename']}")

    def _calculate_priorizacao_payload(self, data: Dict) -> tuple[bool, Dict | str]:
        try:
            qtd_pessoas = float(data.get("qtd_pessoas", 0) or 0)
            horas_mensais = float(data.get("horas_mensais", 0) or 0)
            custo_mensal = float(data.get("custo_mensal", 0) or 0)
            custo_desenvolvimento = float(data.get("custo_desenvolvimento", 0) or 0)
            complexidade = int(data.get("complexidade", 0) or 0)
            dev_interno = int(data.get("dev_interno", 0) or 0)
        except Exception:
            return False, "Entradas invalidas para calculo"

        valores_numericos = (
            qtd_pessoas,
            horas_mensais,
            custo_mensal,
            custo_desenvolvimento,
            float(complexidade),
            float(dev_interno),
        )
        if any(not math.isfinite(valor) for valor in valores_numericos):
            return False, "Os campos numericos devem conter valores finitos."

        if qtd_pessoas <= 0 or custo_mensal <= 0 or horas_mensais <= 0:
            return (
                False,
                "Para calcular, informe qtd_pessoas, custo_mensal e horas_mensais maiores que zero.",
            )
        if custo_desenvolvimento < 0:
            return False, "O custo de desenvolvimento nao pode ser negativo."
        if complexidade < 1 or complexidade > 5:
            return False, "Complexidade deve estar entre 1 e 5."
        if dev_interno < 1 or dev_interno > 5:
            return False, "Desenvolvimento interno deve estar entre 1 e 5."

        custo_hora = custo_mensal / 220
        cp_anual = qtd_pessoas * horas_mensais * custo_hora * 12
        economia_mensal = qtd_pessoas * horas_mensais * custo_hora
        payback_meses = custo_desenvolvimento / economia_mensal if economia_mensal > 0 else float("inf")

        if cp_anual > 600000:
            cp_score = 5
        elif cp_anual > 300000:
            cp_score = 4
        elif cp_anual > 150000:
            cp_score = 3
        elif cp_anual > 50000:
            cp_score = 2
        else:
            cp_score = 1

        if payback_meses <= 3:
            pb_score = 5
        elif payback_meses <= 6:
            pb_score = 4
        elif payback_meses <= 12:
            pb_score = 3
        elif payback_meses <= 24:
            pb_score = 2
        else:
            pb_score = 1

        score_final = round(
            0.35 * cp_score + 0.30 * pb_score + 0.20 * (6 - complexidade) + 0.15 * dev_interno,
            2,
        )

        if score_final >= 4.5:
            prioridade = "Prioridade muito alta"
        elif score_final >= 3.5:
            prioridade = "Prioridade alta"
        elif score_final >= 2.5:
            prioridade = "Prioridade media"
        else:
            prioridade = "Prioridade baixa"

        return True, {
            "custo_hora": round(custo_hora, 2),
            "cp_anual": round(cp_anual, 2),
            "economia_mensal": round(economia_mensal, 2),
            "payback_meses": round(payback_meses, 2),
            "cp_score": cp_score,
            "pb_score": pb_score,
            "score_complexidade": 6 - complexidade,
            "score_dev_interno": dev_interno,
            "score_final": score_final,
            "prioridade": prioridade,
            "metodologia": {
                "premissas": {
                    "horas_base_mes": 220,
                    "pesos": {
                        "cp_score": 0.35,
                        "pb_score": 0.30,
                        "complexidade_invertida": 0.20,
                        "dev_interno": 0.15,
                    },
                },
                "faixas_cp_score": [
                    {"min_exclusive": 600000, "score": 5},
                    {"min_exclusive": 300000, "max_inclusive": 600000, "score": 4},
                    {"min_exclusive": 150000, "max_inclusive": 300000, "score": 3},
                    {"min_exclusive": 50000, "max_inclusive": 150000, "score": 2},
                    {"max_inclusive": 50000, "score": 1},
                ],
                "faixas_pb_score": [
                    {"max_inclusive": 3, "score": 5},
                    {"max_inclusive": 6, "score": 4},
                    {"max_inclusive": 12, "score": 3},
                    {"max_inclusive": 24, "score": 2},
                    {"min_exclusive": 24, "score": 1},
                ],
            },
        }

    def _generate_csrf_token(self) -> str:
        return secrets.token_urlsafe(32)

    def _get_user_permissions(self, user: Dict) -> List[str]:
        roles = self._load_json(self.roles_file, default=[])
        role = next((r for r in roles if r.get("id") == user.get("role_id")), {})
        return list(role.get("permissions") or [])

    def _build_assistant_context_for_user(self, user: Dict, permissions: List[str]) -> Dict[str, Any]:
        perms = set(permissions or [])
        all_access = "*" in perms
        can_access_projects = all_access or ("view-projetos-consultar" in perms) or ("view-projetos-dashboard" in perms)
        can_access_processes = all_access or ("view-processos-consulta" in perms) or ("view-processos-dashboard" in perms) or ("view-processos-cadastro" in perms)
        can_access_registers = all_access or ("view-cadastros-usuarios" in perms) or ("view-cadastros-cargos" in perms) or ("view-projetos-setores" in perms)

        context_parts: List[str] = []

        if can_access_projects:
            projects = self._load_json(self.projects_file, default=[])
            top_projects = projects[:20]
            status_count = self._count_by_key(projects, "status")
            criticidade_count = self._count_by_key(projects, "criticidade")
            responsavel_count = self._count_by_key(projects, "responsavel")
            latest_project_notes = self._collect_project_notes(projects, limit=12)
            context_parts.append(f"Projetos visiveis: {len(projects)}")
            context_parts.append("Projetos por status: " + self._format_count_map(status_count))
            context_parts.append("Projetos por criticidade: " + self._format_count_map(criticidade_count))
            context_parts.append("Projetos por responsavel (top): " + self._format_count_map(responsavel_count, top=8))
            if top_projects:
                context_parts.append(
                    "Resumo projetos: " + " | ".join(
                        [
                            f"{p.get('nome','-')} (status={p.get('status','-')}, resp={p.get('responsavel','-')}, criticidade={p.get('criticidade','-')}, progresso={p.get('progresso',0)}%)"
                            for p in top_projects
                        ]
                    )
                )
            if latest_project_notes:
                context_parts.append(
                    "Anotacoes recentes de projetos: " + " | ".join(
                        [
                            f"[{n.get('data','-')}] {n.get('projeto','-')} - {n.get('autor','-')}: {n.get('texto','-')}"
                            for n in latest_project_notes
                        ]
                    )
                )

        if can_access_processes:
            processos = self._load_json(self.processes_file, default=[])
            top_processos = processos[:20]
            status_count = self._count_by_key(processos, "status")
            depto_count = self._count_by_key(processos, "departamento_nome")
            etapas: List[Dict[str, Any]] = []
            for p in processos:
                etapas.extend(p.get("etapas") or [])
            etapas_status_count = self._count_by_key(etapas, "status")
            etapas_depto_count = self._count_by_key(etapas, "departamento_nome")
            latest_process_comments = self._collect_process_stage_comments(processos, limit=15)
            context_parts.append(f"Processos visiveis: {len(processos)}")
            context_parts.append("Processos por status: " + self._format_count_map(status_count))
            context_parts.append("Processos por departamento: " + self._format_count_map(depto_count))
            context_parts.append("Etapas por status: " + self._format_count_map(etapas_status_count))
            context_parts.append("Etapas por departamento: " + self._format_count_map(etapas_depto_count))
            if top_processos:
                context_parts.append(
                    "Resumo processos: " + " | ".join(
                        [
                            f"{p.get('nome','-')} (status={p.get('status','-')}, depto={p.get('departamento_nome','-')}, resp={p.get('responsavel_nome','-')}, etapas={len(p.get('etapas') or [])})"
                            for p in top_processos
                        ]
                    )
                )
            if latest_process_comments:
                context_parts.append(
                    "Comentarios recentes de etapas: " + " | ".join(
                        [
                            f"[{c.get('data','-')}] {c.get('processo','-')} / {c.get('etapa','-')} - {c.get('autor','-')}: {c.get('texto','-')}"
                            for c in latest_process_comments
                        ]
                    )
                )

        if can_access_registers:
            departamentos = self._load_json(self.sectors_file, default=[])
            cargos = self._load_json(self.cargos_file, default=[])
            setores = self._load_json(self.sectors_file, default=[])
            users = self._load_json(self.users_file, default=[])
            context_parts.append(
                f"Cadastros visiveis: departamentos={len(departamentos)}, cargos={len(cargos)}, setores={len(setores)}, usuarios={len(users)}"
            )

        if not context_parts:
            context_parts.append("Sem acesso a dados de projetos, processos ou cadastros para consulta.")

        return {
            "user_name": user.get("nome", "Usuario"),
            "permissions": sorted(list(perms)),
            "can_access_projects": can_access_projects,
            "can_access_processes": can_access_processes,
            "can_access_registers": can_access_registers,
            "context_text": "\n".join(context_parts),
        }

    def _assistant_local_fallback_answer(self, message: str, context_payload: Dict[str, Any]) -> str:
        normalized = self._normalize_text(message)
        answer = self._assistant_answer_process_count_by_department(normalized, context_payload)
        if answer:
            return answer
        answer = self._assistant_answer_process_comments(normalized, context_payload)
        if answer:
            return answer
        answer = self._assistant_answer_project_annotations(normalized, context_payload)
        if answer:
            return answer
        return (
            "Ola! Eu sou a Justine, assistente do HPO. No momento estou em modo local de contingencia. "
            "Posso responder consultas de quantidade por area/departamento/status e comentarios/anotacoes com base nos seus acessos."
        )

    def _assistant_answer_process_count_by_department(self, normalized_message: str, context_payload: Dict[str, Any]) -> str:
        if not context_payload.get("can_access_processes"):
            return "Voce nao possui permissao para consultar dados de processos."

        count_keywords = ("quantos", "quantidade", "qtd", "numero", "total")
        process_keywords = ("processo", "processos")
        is_count_question = any(k in normalized_message for k in count_keywords)
        mentions_process = any(k in normalized_message for k in process_keywords)

        if not (is_count_question and mentions_process):
            return ""

        processos = self._load_json(self.processes_file, default=[])
        depto_map: Dict[str, int] = {}
        for proc in processos:
            depto_name = str(proc.get("departamento_nome") or proc.get("departamento") or "").strip()
            if not depto_name:
                continue
            key = self._normalize_text(depto_name)
            depto_map[key] = depto_map.get(key, 0) + 1

        if not depto_map:
            return "Nao encontrei processos cadastrados para consulta."

        # Primeiro tenta identificar departamento citado livremente na frase
        queried_department = self._match_department_in_message(normalized_message, list(depto_map.keys()))
        if not queried_department:
            queried_department = self._extract_department_name(normalized_message)
        if queried_department:
            for depto_key, total in depto_map.items():
                if queried_department in depto_key or depto_key in queried_department:
                    return f"Atualmente temos {total} processo(s) no departamento de {depto_key.title()}."
            return "Nao encontrei departamento correspondente na base para essa consulta."

        ordered = sorted(depto_map.items(), key=lambda item: item[1], reverse=True)
        top = " | ".join([f"{name.title()}: {qty}" for name, qty in ordered[:5]])
        return f"Total de processos por departamento (top 5): {top}."

    def _extract_department_name(self, normalized_message: str) -> str:
        match = re.search(r"\b(?:no|na|do|da|de)\s+([a-z0-9\s_-]{2,})[\.\,\!\?\:\;]*$", normalized_message)
        if not match:
            return ""
        candidate = match.group(1).strip(" .,!?:;")
        return candidate

    def _match_department_in_message(self, normalized_message: str, department_keys: List[str]) -> str:
        # Match por ocorrencia direta de nome de departamento na frase inteira.
        # Prioriza nomes maiores para evitar ambiguidades.
        for depto in sorted((d for d in department_keys if d), key=len, reverse=True):
            if re.search(rf"\b{re.escape(depto)}\b", normalized_message):
                return depto
        return ""

    def _assistant_answer_process_comments(self, normalized_message: str, context_payload: Dict[str, Any]) -> str:
        if not context_payload.get("can_access_processes"):
            return ""
        if not any(k in normalized_message for k in ("comentario", "comentarios")):
            return ""

        processos = self._load_json(self.processes_file, default=[])
        if not processos:
            return "Nao encontrei processos cadastrados para consulta de comentarios."

        depto_names = sorted(
            {self._normalize_text(str(p.get("departamento_nome", "")).strip()) for p in processos if p.get("departamento_nome")}
        )
        depto_in_msg = self._match_department_in_message(normalized_message, depto_names)

        comments = self._collect_process_stage_comments(processos, limit=200)
        if depto_in_msg:
            comments = [c for c in comments if self._normalize_text(c.get("departamento", "")) == depto_in_msg]
        if not comments:
            return "Nao encontrei comentarios de etapas para o filtro solicitado."

        top = comments[:5]
        lines = [
            f"- [{c.get('data','-')}] {c.get('processo','-')} / {c.get('etapa','-')} - {c.get('autor','-')}: {c.get('texto','-')}"
            for c in top
        ]
        suffix = f" no departamento de {depto_in_msg.title()}" if depto_in_msg else ""
        return f"Encontrei {len(comments)} comentario(s) de etapas{suffix}. Mais recentes:\n" + "\n".join(lines)

    def _assistant_answer_project_annotations(self, normalized_message: str, context_payload: Dict[str, Any]) -> str:
        if not context_payload.get("can_access_projects"):
            return ""
        if not any(k in normalized_message for k in ("anotacao", "anotacoes", "comentario", "comentarios")):
            return ""
        if "projeto" not in normalized_message and "projetos" not in normalized_message:
            return ""

        projects = self._load_json(self.projects_file, default=[])
        notes = self._collect_project_notes(projects, limit=200)
        if not notes:
            return "Nao encontrei anotacoes de projetos para consulta."
        top = notes[:5]
        lines = [
            f"- [{n.get('data','-')}] {n.get('projeto','-')} - {n.get('autor','-')}: {n.get('texto','-')}"
            for n in top
        ]
        return f"Encontrei {len(notes)} anotacao(oes) em projetos. Mais recentes:\n" + "\n".join(lines)

    def _count_by_key(self, items: List[Dict[str, Any]], key: str) -> Dict[str, int]:
        count: Dict[str, int] = {}
        for item in items:
            value = str(item.get(key, "")).strip() or "nao_informado"
            count[value] = count.get(value, 0) + 1
        return count

    def _can_user_access_project(self, project: Dict[str, Any], user: Optional[Dict[str, Any]]) -> bool:
        if not user:
            return False
        user_id = str(user.get("id", "")).strip()
        user_nome = str(user.get("nome", "")).strip().lower()
        username = str(user.get("username", "")).strip().lower()
        permissions = set(self._get_user_permissions(user) or [])
        if username == "admin" or "*" in permissions:
            return True

        responsavel_id = str(project.get("responsavel_id", "")).strip()
        responsavel_nome = str(project.get("responsavel", "")).strip().lower()
        participantes_ids = {str(pid).strip() for pid in (project.get("participantes_ids") or []) if str(pid).strip()}
        participantes_nomes = {
            str(nome).strip().lower()
            for nome in ((project.get("participantes") or []) + (project.get("setores_impactados") or []))
            if str(nome).strip()
        }

        if user_id and responsavel_id == user_id:
            return True
        if user_nome and responsavel_nome == user_nome:
            return True
        if user_id and user_id in participantes_ids:
            return True
        if user_nome and user_nome in participantes_nomes:
            return True
        for etapa in project.get("etapas") or []:
            etapa_responsavel_nome = str(etapa.get("responsavel", "")).strip().lower()
            if user_nome and etapa_responsavel_nome == user_nome:
                return True
        return False

    def _filter_projects_for_user(self, projects: List[Dict[str, Any]], user: Optional[Dict[str, Any]]) -> List[Dict[str, Any]]:
        return [project for project in (projects or []) if self._can_user_access_project(project, user)]

    def _format_count_map(self, count_map: Dict[str, int], top: int = 12) -> str:
        if not count_map:
            return "sem dados"
        ordered = sorted(count_map.items(), key=lambda kv: kv[1], reverse=True)
        return " | ".join([f"{k}: {v}" for k, v in ordered[:top]])

    def _collect_project_notes(self, projects: List[Dict[str, Any]], limit: int = 12) -> List[Dict[str, str]]:
        notes: List[Dict[str, str]] = []
        for p in projects:
            projeto_nome = str(p.get("nome", "")).strip() or "-"
            for n in (p.get("anotacoes") or []):
                texto = str(n.get("conteudo", "")).strip()
                if not texto:
                    continue
                notes.append(
                    {
                        "projeto": projeto_nome,
                        "data": str(n.get("data", "")).strip() or "-",
                        "autor": str(n.get("usuario", "")).strip() or "-",
                        "texto": texto,
                    }
                )
            resumo = str(p.get("anotacoes_gerais", "")).strip()
            if resumo:
                notes.append(
                    {
                        "projeto": projeto_nome,
                        "data": str(p.get("ultima_atualizacao", "")).strip() or str(p.get("created_at", "")).strip() or "-",
                        "autor": str(p.get("responsavel", "")).strip() or "-",
                        "texto": resumo,
                    }
                )
        notes.sort(key=lambda n: str(n.get("data", "")), reverse=True)
        return notes[:limit]

    def _collect_process_stage_comments(self, processos: List[Dict[str, Any]], limit: int = 15) -> List[Dict[str, str]]:
        comments: List[Dict[str, str]] = []
        for p in processos:
            processo_nome = str(p.get("nome", "")).strip() or "-"
            for e in (p.get("etapas") or []):
                etapa_nome = str(e.get("nome", "")).strip() or "-"
                etapa_depto = str(e.get("departamento_nome", "")).strip() or str(p.get("departamento_nome", "")).strip() or "-"
                for c in ((e.get("conhecimento") or {}).get("comentarios") or []):
                    texto = str(c.get("texto", "")).strip()
                    if not texto:
                        continue
                    comments.append(
                        {
                            "processo": processo_nome,
                            "etapa": etapa_nome,
                            "departamento": etapa_depto,
                            "data": str(c.get("data_hora", "")).strip() or "-",
                            "autor": str(c.get("autor", "")).strip() or "-",
                            "texto": texto,
                        }
                    )
        comments.sort(key=lambda c: str(c.get("data", "")), reverse=True)
        return comments[:limit]

    def _normalize_text(self, text: str) -> str:
        base = unicodedata.normalize("NFKD", str(text or ""))
        ascii_text = "".join(ch for ch in base if not unicodedata.combining(ch))
        ascii_text = ascii_text.lower().strip()
        ascii_text = re.sub(r"\s+", " ", ascii_text)
        return ascii_text

    def _should_validate_csrf(self, path: str, method: str) -> bool:
        if not self.enable_csrf:
            return False
        if method in {"GET", "HEAD", "OPTIONS"}:
            return False
        if not path.startswith("/api/"):
            return False
        return path not in self.csrf_exempt_paths

    def _client_ip(self) -> str:
        forwarded_for = request.headers.get("X-Forwarded-For", "").strip()
        if forwarded_for:
            return forwarded_for.split(",")[0].strip()
        return request.remote_addr or "unknown"

    def _rate_limit_key(self, kind: str, value: str) -> str:
        return f"{kind}:{value or 'unknown'}"

    def _prune_rate_limit(self, key: str, now_ts: float) -> List[float]:
        window_start = now_ts - self.login_rate_window_sec
        entries = self.login_rate_limit_store.get(key, [])
        filtered = [ts for ts in entries if ts >= window_start]
        self.login_rate_limit_store[key] = filtered
        return filtered

    def _check_login_rate_limit(self, ip: str, username: str) -> tuple[bool, str]:
        with self.login_rate_lock:
            now_ts = time.time()
            ip_entries = self._prune_rate_limit(self._rate_limit_key("ip", ip), now_ts)
            user_entries = self._prune_rate_limit(self._rate_limit_key("user", username), now_ts)
            if len(ip_entries) >= self.login_rate_ip_max:
                return False, "Muitas tentativas de login. Tente novamente em alguns minutos."
            if len(user_entries) >= self.login_rate_user_max:
                return False, "Muitas tentativas para este usuario. Tente novamente em alguns minutos."
            return True, ""

    def _register_login_failure(self, ip: str, username: str):
        with self.login_rate_lock:
            now_ts = time.time()
            ip_key = self._rate_limit_key("ip", ip)
            user_key = self._rate_limit_key("user", username)
            self.login_rate_limit_store.setdefault(ip_key, []).append(now_ts)
            self.login_rate_limit_store.setdefault(user_key, []).append(now_ts)

    def _clear_login_rate_limit(self, ip: str, username: str):
        with self.login_rate_lock:
            self.login_rate_limit_store.pop(self._rate_limit_key("ip", ip), None)
            self.login_rate_limit_store.pop(self._rate_limit_key("user", username), None)

    def _is_temp_password_expired(self, user: Dict) -> bool:
        if not bool(user.get("password_is_temporary", False)):
            return False
        expires_at = str(user.get("temp_password_expires_at", "")).strip()
        if not expires_at:
            return False
        try:
            return datetime.fromisoformat(expires_at) < datetime.now()
        except Exception:
            return False

    def _generate_temporary_password(self, length: int = 12) -> str:
        alphabet = "ABCDEFGHJKLMNPQRSTUVWXYZabcdefghijkmnopqrstuvwxyz23456789!@#$%"
        return "".join(secrets.choice(alphabet) for _ in range(max(10, length)))

    def _build_repository_store(self):
        use_supabase = str(os.getenv("SUPABASE_ENABLED", "0")).strip() == "1"
        if not use_supabase:
            return JsonRepositoryStore()

        supabase_url = str(os.getenv("SUPABASE_URL", "")).strip()
        supabase_service_key = str(os.getenv("SUPABASE_SERVICE_ROLE_KEY", "")).strip()
        table_map = {
            "projects.json": "hpo_projects",
            "processes_registry.json": "hpo_processes",
            "users.json": "hpo_users",
            "roles.json": "hpo_roles",
            "departments.json": "hpo_departments",
            "cargos.json": "hpo_cargos",
            "sectors.json": "hpo_sectors",
        }
        try:
            store = SupabaseRepositoryStore(
                url=supabase_url,
                service_role_key=supabase_service_key,
                table_map=table_map,
            )
            self.logger.info("Persistencia Supabase habilitada via SUPABASE_ENABLED=1.")
            return store
        except Exception as exc:
            self.logger.warning(f"Falha ao habilitar Supabase. Fallback para JSON local. Motivo: {exc}")
            return JsonRepositoryStore()

    def _init_data_store(self):
        self.data_dir.mkdir(parents=True, exist_ok=True)
        if not self.projects_file.exists():
            self._save_json(self.projects_file, [])
        if not self.sectors_file.exists():
            self._save_json(self.sectors_file, [])
        if not self.departments_file.exists():
            self._save_json(self.departments_file, [])
        if not self.cargos_file.exists():
            self._save_json(self.cargos_file, [])
        if not self.processes_file.exists():
            self._save_json(self.processes_file, [])
        if not self.roles_file.exists():
            admin_role = {"id": str(uuid4()), "nome": "Administrador", "permissions": ["*"]}
            self._save_json(self.roles_file, [admin_role])
        if not self.users_file.exists():
            roles = self._load_json(self.roles_file, default=[])
            admin_role = next((r for r in roles if r.get("nome") == "Administrador"), None)
            admin_role_id = admin_role.get("id") if admin_role else ""
            admin_initial_password = str(os.getenv("ADMIN_INITIAL_PASSWORD", "")).strip()
            admin_password_is_temporary = False
            admin_password_expires_at = ""
            admin_must_change_password = False
            if len(admin_initial_password) < 10:
                admin_initial_password = self._generate_temporary_password()
                admin_password_is_temporary = True
                admin_must_change_password = True
                admin_password_expires_at = (
                    datetime.now() + timedelta(hours=int(os.getenv("TEMP_PASSWORD_EXP_HOURS", "24")))
                ).isoformat()
                bootstrap_file = self.data_dir / "admin_bootstrap_password.txt"
                bootstrap_file.write_text(
                    f"Senha temporaria inicial do admin: {admin_initial_password}\n"
                    f"Expira em: {admin_password_expires_at}\n",
                    encoding="utf-8",
                )
                self.logger.warning("ADMIN_INITIAL_PASSWORD nao definido. Senha inicial temporaria gerada.")
            admin_user = {
                "id": str(uuid4()),
                "nome": "Administrador",
                "email": "admin@hpo.local",
                "username": "admin",
                "role_id": admin_role_id,
                "departamento_id": "",
                "cargo_id": "",
                "has_login": True,
                "password_hash": self._hash_password(admin_initial_password),
                "must_change_password": admin_must_change_password,
                "password_is_temporary": admin_password_is_temporary,
                "temp_password_expires_at": admin_password_expires_at,
                "active": True,
                "created_at": datetime.now().isoformat(),
                "updated_at": datetime.now().isoformat(),
            }
            self._save_json(self.users_file, [admin_user])
        else:
            users = self._load_json(self.users_file, default=[])
            changed = False
            for u in users:
                if "has_login" not in u:
                    u["has_login"] = True
                    changed = True
                if "departamento_id" not in u:
                    u["departamento_id"] = ""
                    changed = True
                if "cargo_id" not in u:
                    u["cargo_id"] = ""
                    changed = True
                if "password_is_temporary" not in u:
                    u["password_is_temporary"] = False
                    changed = True
                if "temp_password_expires_at" not in u:
                    u["temp_password_expires_at"] = ""
                    changed = True
            if changed:
                self._save_json(self.users_file, users)

    def _load_json(self, file_path: Path, default):
        try:
            data = self.repo_store.read(file_path, default)
            normalized = self._normalize_text_tree(data)
            if normalized != data:
                self._save_json(file_path, normalized)
            return normalized
        except Exception:
            return default

    def _save_json(self, file_path: Path, data):
        normalized = self._normalize_text_tree(data)
        self.repo_store.write(file_path, normalized)

    def _normalize_text_tree(self, value):
        if isinstance(value, dict):
            return {k: self._normalize_text_tree(v) for k, v in value.items()}
        if isinstance(value, list):
            return [self._normalize_text_tree(v) for v in value]
        if isinstance(value, str):
            return self._fix_mojibake(value)
        return value

    def _fix_mojibake(self, text: str) -> str:
        # Heuristic fix for strings previously decoded with a wrong codepage (e.g. "FÃ¡bio").
        try:
            bad_patterns = (
                "Ã¡", "Ã¢", "Ã£", "Ã¤", "Ã©", "Ãª", "Ã­", "Ã³", "Ã´", "Ãµ", "Ãº", "Ã§",
                "Ã ", "Ã‰", "Ã“", "Ã‡", "Â", "â€“", "â€”", "â€œ", "â€", "�"
            )
            if any(token in text for token in bad_patterns):
                repaired = text.encode("latin1", errors="ignore").decode("utf-8", errors="ignore")
                return repaired or text
        except Exception:
            pass
        return text

    def _hash_password(self, password: str) -> str:
        return hashlib.sha256(f"{self.auth_salt}:{password}".encode("utf-8")).hexdigest()

    def _serialize_user_public(self, user: Dict) -> Dict:
        return {
            "id": user.get("id"),
            "nome": user.get("nome"),
            "email": user.get("email"),
            "username": user.get("username"),
            "role_id": user.get("role_id"),
            "departamento_id": user.get("departamento_id", ""),
            "cargo_id": user.get("cargo_id", ""),
            "has_login": bool(user.get("has_login", True)),
            "active": bool(user.get("active", True)),
            "must_change_password": bool(user.get("must_change_password", False)),
        }

    def _serialize_user_session(self, user: Dict) -> Dict:
        roles = self._load_json(self.roles_file, default=[])
        departamentos = self._load_json(self.sectors_file, default=[])
        cargos = self._load_json(self.cargos_file, default=[])
        role = next((r for r in roles if r.get("id") == user.get("role_id")), {})
        dept = next((d for d in departamentos if d.get("id") == user.get("departamento_id")), {})
        cargo = next((c for c in cargos if c.get("id") == user.get("cargo_id")), {})
        permissions = role.get("permissions") or []
        return {
            **self._serialize_user_public(user),
            "role_nome": role.get("nome", ""),
            "departamento_nome": dept.get("nome", ""),
            "cargo_nome": cargo.get("nome", ""),
            "permissions": permissions,
        }

    def _get_current_user(self) -> Optional[Dict]:
        user_id = session.get("user_id")
        if not user_id:
            return None
        users = self._load_json(self.users_file, default=[])
        return next((u for u in users if u.get("id") == user_id), None)

    def _require_permission(self, view_id: str):
        user = self._get_current_user()
        if not user:
            return jsonify({"ok": False, "message": "Nao autenticado"}), 401
        roles = self._load_json(self.roles_file, default=[])
        role = next((r for r in roles if r.get("id") == user.get("role_id")), {})
        perms = role.get("permissions") or []
        if "*" in perms or view_id in perms:
            return None
        return jsonify({"ok": False, "message": "Sem permissao"}), 403

    def _validate_project_payload(self, payload: Dict) -> tuple[bool, Dict | str]:
        nome = str(payload.get("nome", "")).strip()
        descricao = str(payload.get("descricao", "")).strip()
        responsavel = str(payload.get("responsavel", "")).strip()
        responsavel_id = str(payload.get("responsavel_id", "")).strip()
        focal = str(payload.get("focal", "")).strip()
        focal_id = str(payload.get("focal_id", "")).strip()
        users = self._load_json(self.users_file, default=[])
        responsavel_user = next((u for u in users if str(u.get("id", "")).strip() == responsavel_id), None) if responsavel_id else None
        focal_user = next((u for u in users if str(u.get("id", "")).strip() == focal_id), None) if focal_id else None
        if responsavel_id and not responsavel_user:
            return False, "Responsável do projeto inválido"
        if focal_id and not focal_user:
            return False, "Focal do projeto inválido"
        if responsavel_user:
            responsavel = str(responsavel_user.get("nome") or responsavel_user.get("email") or responsavel_user.get("username") or "").strip()
        elif responsavel:
            responsavel_user = next((u for u in users if str(u.get("nome", "")).strip().lower() == responsavel.lower()), None)
            responsavel_id = str((responsavel_user or {}).get("id", "")).strip()
        if focal_user:
            focal = str(focal_user.get("nome") or focal_user.get("email") or focal_user.get("username") or "").strip()
        elif focal:
            focal_user = next((u for u in users if str(u.get("nome", "")).strip().lower() == focal.lower()), None)
            focal_id = str((focal_user or {}).get("id", "")).strip()

        if not nome or not descricao or not responsavel:
            return False, "Nome, descrição e responsável são obrigatórios"

        participantes_ids_raw = payload.get("participantes_ids")
        if participantes_ids_raw is None:
            participantes_ids_raw = payload.get("participantes")
        if participantes_ids_raw is None:
            participantes_ids_raw = payload.get("setores_impactados")
        participantes_ids = []
        participantes_nomes = []
        for raw_participante in participantes_ids_raw or []:
            participante_id = str(raw_participante).strip()
            if not participante_id:
                continue
            participante_user = next((u for u in users if str(u.get("id", "")).strip() == participante_id), None)
            if participante_user:
                participante_nome = str(
                    participante_user.get("nome") or participante_user.get("email") or participante_user.get("username") or ""
                ).strip()
                if participante_id not in participantes_ids:
                    participantes_ids.append(participante_id)
                if participante_nome and participante_nome not in participantes_nomes:
                    participantes_nomes.append(participante_nome)
            else:
                if participante_id not in participantes_nomes:
                    participantes_nomes.append(participante_id)

        data_inicio_previsto = str(payload.get("data_inicio_previsto", "")).strip()
        dt_inicio_real = str(payload.get("dt_inicio_real", "")).strip()
        previsao_termino = str(payload.get("previsao_termino", "")).strip()
        termino_real = str(payload.get("termino_real", "")).strip()

        date_checks = [
            ("Data início previsto", data_inicio_previsto),
            ("Data início real", dt_inicio_real),
            ("Previsão término", previsao_termino),
            ("Término real", termino_real),
        ]
        for label, raw_date in date_checks:
            if raw_date and not self._parse_date_yyyy_mm_dd(raw_date):
                return False, f"{label} inválida. Use formato YYYY-MM-DD com data real."

        ini_prev_dt = self._parse_date_yyyy_mm_dd(data_inicio_previsto)
        fim_prev_dt = self._parse_date_yyyy_mm_dd(previsao_termino)
        if ini_prev_dt and fim_prev_dt and fim_prev_dt < ini_prev_dt:
            return False, "Previsão de término não pode ser menor que a data de início previsto."

        ini_real_dt = self._parse_date_yyyy_mm_dd(dt_inicio_real)
        fim_real_dt = self._parse_date_yyyy_mm_dd(termino_real)
        if ini_real_dt and fim_real_dt and fim_real_dt < ini_real_dt:
            return False, "Término real não pode ser menor que a data de início real."

        etapas = payload.get("etapas") or []
        etapas_norm = []
        for etapa in etapas:
            fim_previsto = str(etapa.get("fim_previsto", "")).strip()
            prazo = str(etapa.get("prazo", "")).strip() or fim_previsto
            etapa_nome = str(etapa.get("nome", "")).strip() or "Etapa sem nome"
            inicio_previsto = str(etapa.get("inicio_previsto", "")).strip()
            inicio_real = str(etapa.get("inicio_real", "")).strip()
            fim_real = str(etapa.get("fim_real", "")).strip()

            etapa_date_checks = [
                ("início previsto", inicio_previsto),
                ("início real", inicio_real),
                ("fim previsto", fim_previsto),
                ("fim real", fim_real),
                ("prazo", prazo),
            ]
            for label, raw_date in etapa_date_checks:
                if raw_date and not self._parse_date_yyyy_mm_dd(raw_date):
                    return False, f"{etapa_nome}: {label} inválido(a). Use formato YYYY-MM-DD com data real."

            etapa_ini_prev_dt = self._parse_date_yyyy_mm_dd(inicio_previsto)
            etapa_fim_prev_dt = self._parse_date_yyyy_mm_dd(fim_previsto)
            etapa_ini_real_dt = self._parse_date_yyyy_mm_dd(inicio_real)
            etapa_fim_real_dt = self._parse_date_yyyy_mm_dd(fim_real)
            etapa_prazo_dt = self._parse_date_yyyy_mm_dd(prazo)

            if etapa_ini_prev_dt and etapa_fim_prev_dt and etapa_fim_prev_dt < etapa_ini_prev_dt:
                return False, f"{etapa_nome}: fim previsto não pode ser menor que início previsto."
            if etapa_ini_real_dt and etapa_fim_real_dt and etapa_fim_real_dt < etapa_ini_real_dt:
                return False, f"{etapa_nome}: fim real não pode ser menor que início real."
            if etapa_ini_prev_dt and etapa_prazo_dt and etapa_prazo_dt < etapa_ini_prev_dt:
                return False, f"{etapa_nome}: prazo não pode ser menor que início previsto."

            etapas_norm.append(
                {
                    "id": etapa.get("id") or str(uuid4()),
                    "nome": str(etapa.get("nome", "")).strip(),
                    "responsavel": str(etapa.get("responsavel", "")).strip(),
                    "inicio_previsto": inicio_previsto,
                    "inicio_real": inicio_real,
                    "fim_previsto": fim_previsto or prazo,
                    "fim_real": fim_real,
                    "prazo": prazo,
                    "descricao": str(etapa.get("descricao", "")).strip(),
                    "anotacoes": str(etapa.get("anotacoes", "")).strip(),
                    "criticidade": str(etapa.get("criticidade", "")).strip(),
                    "complexidade": str(etapa.get("complexidade", "")).strip(),
                    "status": str(etapa.get("status", "pendente")).strip() or "pendente",
                    "anexos": etapa.get("anexos") or [],
                }
            )
        etapas_norm.sort(key=lambda e: e.get("prazo") or "9999-12-31")

        progresso = self._compute_progress(etapas_norm, payload.get("status"))

        project = {
            "nome": nome,
            "descricao": descricao,
            "responsavel": responsavel,
            "responsavel_id": responsavel_id,
            "focal": focal,
            "focal_id": focal_id,
            "status": str(payload.get("status", "backlog")).strip(),
            "criticidade": str(payload.get("criticidade", "media")).strip(),
            "data_inicio_previsto": data_inicio_previsto,
            "dt_inicio_real": dt_inicio_real,
            "previsao_termino": previsao_termino,
            "termino_real": termino_real,
            "setor_projeto": str(payload.get("setor_projeto", "")).strip(),
            "participantes_ids": participantes_ids,
            "participantes": participantes_nomes,
            "setores_impactados": participantes_nomes,
            "anexos": payload.get("anexos") or [],
            "etapas": etapas_norm,
            "anotacoes_gerais": str(payload.get("anotacoes_gerais", "")).strip(),
            "anotacoes": payload.get("anotacoes") or [],
            "progresso": progresso,
            "tarefas": len(etapas_norm),
            "tarefasConcluidas": sum(1 for e in etapas_norm if e.get("status") == "concluida"),
            "orcamento": float(payload.get("orcamento", 0) or 0),
            "custo_atual": float(payload.get("custo_atual", 0) or 0),
            "setor": str(payload.get("setor_projeto", "")).strip(),
        }
        return True, project

    def _validate_processo_payload(self, payload: Dict, processo_id: Optional[str]) -> tuple[bool, Dict | str]:
        nome = str(payload.get("nome", "")).strip()
        descricao = str(payload.get("descricao", "")).strip()
        status = str(payload.get("status", "rascunho")).strip().lower()
        versao = str(payload.get("versao", "1.0")).strip() or "1.0"
        departamento_id = str(payload.get("departamento_id", "")).strip()
        responsavel_id = str(payload.get("responsavel_id", "")).strip()
        if not nome or not descricao or not departamento_id or not responsavel_id:
            return False, "Nome, descricao, departamento e responsavel sao obrigatorios"
        if status not in {"ativo", "inativo", "rascunho"}:
            return False, "Status invalido para processo"

        departamentos = self._load_json(self.sectors_file, default=[])
        cargos = self._load_json(self.cargos_file, default=[])
        users = self._load_json(self.users_file, default=[])
        dept = next((d for d in departamentos if d.get("id") == departamento_id), None)
        if not dept:
            return False, "Departamento do processo invalido"
        responsavel = next((u for u in users if u.get("id") == responsavel_id), None)
        if not responsavel:
            return False, "Responsavel do processo invalido"

        etapas = payload.get("etapas") or []
        etapas_norm: List[Dict[str, Any]] = []
        for raw in etapas:
            etapa_nome = str(raw.get("nome", "")).strip()
            if not etapa_nome:
                return False, "Nome da etapa e obrigatorio"
            etapa_departamento_id = str(raw.get("departamento_id", "")).strip() or departamento_id
            etapa_dept = next((d for d in departamentos if d.get("id") == etapa_departamento_id), None)
            if not etapa_dept:
                return False, f"Departamento invalido na etapa: {etapa_nome}"
            etapa_cargo_id = str(raw.get("cargo_id", "")).strip()
            etapa_cargo = next((c for c in cargos if c.get("id") == etapa_cargo_id), None) if etapa_cargo_id else None
            if etapa_cargo_id and not etapa_cargo:
                return False, f"Cargo invalido na etapa: {etapa_nome}"
            etapa_responsavel_id = str(raw.get("responsavel_id", "")).strip()
            etapa_responsavel = next((u for u in users if u.get("id") == etapa_responsavel_id), None) if etapa_responsavel_id else None
            if etapa_responsavel_id and not etapa_responsavel:
                return False, f"Responsavel direto invalido na etapa: {etapa_nome}"
            ordem = int(raw.get("ordem", 0) or 0)
            if ordem <= 0:
                ordem = len(etapas_norm) + 1

            conhecimento = raw.get("conhecimento") or {}
            anexos = conhecimento.get("anexos") or []
            checklist = conhecimento.get("checklist") or []
            comentarios = conhecimento.get("comentarios") or []
            comentarios_norm = []
            for com in comentarios:
                comentarios_norm.append(
                    {
                        "id": com.get("id") or str(uuid4()),
                        "autor_id": str(com.get("autor_id", "")).strip(),
                        "autor_nome": str(com.get("autor_nome", "")).strip(),
                        "data_hora": str(com.get("data_hora", datetime.now().isoformat())).strip(),
                        "texto": str(com.get("texto", "")).strip(),
                        "edit_history": com.get("edit_history") or [],
                    }
                )
            comentarios_norm.sort(key=lambda c: str(c.get("data_hora", "")))
            checklist_norm = [
                {
                    "id": item.get("id") or str(uuid4()),
                    "texto": str(item.get("texto", "")).strip(),
                    "concluido": bool(item.get("concluido", False)),
                }
                for item in checklist
                if str(item.get("texto", "")).strip()
            ]

            etapas_norm.append(
                {
                    "id": raw.get("id") or str(uuid4()),
                    "processo_id": processo_id or "",
                    "nome": etapa_nome,
                    "ordem": ordem,
                    "departamento_id": etapa_departamento_id,
                    "departamento_nome": etapa_dept.get("nome", ""),
                    "cargo_id": etapa_cargo_id,
                    "cargo_nome": (etapa_cargo or {}).get("nome", ""),
                    "responsavel_id": etapa_responsavel_id,
                    "responsavel_nome": (etapa_responsavel or {}).get("nome", ""),
                    "descricao": str(raw.get("descricao", "")).strip(),
                    "sla": str(raw.get("sla", "")).strip(),
                    "tipo_entrada": str(raw.get("tipo_entrada", "manual")).strip().lower() or "manual",
                    "status": str(raw.get("status", "ativa")).strip().lower() or "ativa",
                    "conhecimento": {
                        "instrucoes": str(conhecimento.get("instrucoes", "")).strip(),
                        "observacoes": str(conhecimento.get("observacoes", "")).strip(),
                        "pontos_atencao": str(conhecimento.get("pontos_atencao", "")).strip(),
                        "checklist": checklist_norm,
                        "anexos": anexos,
                        "comentarios": comentarios_norm,
                    },
                }
            )

        etapas_norm.sort(key=lambda e: int(e.get("ordem", 0)))
        for idx, etapa in enumerate(etapas_norm):
            etapa["ordem"] = idx + 1
            etapa["processo_id"] = processo_id or ""

        return True, {
            "nome": nome,
            "departamento_id": departamento_id,
            "departamento_nome": dept.get("nome", ""),
            "descricao": descricao,
            "status": status,
            "versao": versao,
            "responsavel_id": responsavel_id,
            "responsavel_nome": responsavel.get("nome", ""),
            "etapas": etapas_norm,
        }

    def _parse_date_yyyy_mm_dd(self, date_str: str) -> Optional[datetime]:
        try:
            return datetime.strptime(date_str, "%Y-%m-%d")
        except Exception:
            return None

    def _compute_progress(self, etapas: List[Dict], project_status: Optional[str] = None) -> float:
        if not etapas:
            return 100.0 if project_status == "concluido" else 0.0
        total = len(etapas)
        concluidas = sum(1 for e in etapas if e.get("status") == "concluida")
        return round((concluidas / total) * 100, 2)

